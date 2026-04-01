# -*- coding: utf-8 -*-
"""
휙 단지 조회 + 원고 생성 엔진 (hwik_engine)
- Claude AI 자유형식 입력 파싱
- 도로명 주소 우선 검색
- 매매/전세/월세 실거래가
- 주변 단지 실거래가
- 카카오 학교 정보 (SC4)
- 이미지 4장 생성 (단순지도 / 매매라벨 / 전세라벨 / 그래프)
- docx 원고 8개 생성 (기존 포맷 동일)
"""

import requests
import urllib3
import ssl
import xml.etree.ElementTree as ET
import json
import math
import time
import os
import sys
import random
import argparse
from io import BytesIO
from datetime import datetime
from dateutil.relativedelta import relativedelta
from requests.adapters import HTTPAdapter
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
import re
from collections import Counter

urllib3.disable_warnings()

# ===== docx =====
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 그래프 =====
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import matplotlib.ticker as mticker

# ===== 이미지 합성 =====
from PIL import Image, ImageDraw, ImageFont


# ========================================================
# 설정 (config.py에서 API 키 로드)
# ========================================================
from config import load_env, GOV_SERVICE_KEY, KAKAO_API_KEY, KAKAO_JS_KEY, ANTHROPIC_API_KEY, ANTHROPIC_MODEL_HAIKU, anthropic_headers

UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
HDR_XML  = {"Accept": "application/xml",  "User-Agent": UA}
HDR_JSON = {"Accept": "application/json", "User-Agent": UA}

SUPABASE_URL     = os.environ.get('SUPABASE_URL', 'https://api.hwik.kr')
SUPABASE_KEY     = os.environ.get('SUPABASE_KEY', '')
SUPABASE_HEADERS = {
    "apikey":        SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
}

# ── danji_pages DB에서 캐시 데이터 조회 ──
def fetch_danji_page_data(apt_name, gu="", dong=""):
    """danji_pages 테이블에서 단지 데이터 조회 (블로그 원고용 캐시)"""
    if not SUPABASE_KEY:
        return None
    try:
        # 단지명으로 검색
        search = apt_name.replace(" ", "")
        r = requests.get(
            f"{SUPABASE_URL}/rest/v1/danji_pages",
            headers=SUPABASE_HEADERS,
            params={
                "select": "id,complex_name,location,address,build_year,total_units,categories,recent_trade,all_time_high,jeonse_rate,price_history,nearby_subway,nearby_school,nearby_complex,lat,lng,top_floor,parking,heating,builder,pyeongs_map",
                "complex_name": f"ilike.*{search}*",
                "limit": "5",
            },
            timeout=15,
        )
        if r.status_code != 200 or not r.json():
            return None
        candidates = r.json()
        # 구/동으로 필터
        best = None
        for c in candidates:
            loc = c.get("location", "")
            if gu and gu not in loc:
                continue
            if dong and dong not in loc:
                continue
            best = c
            break
        if not best and candidates:
            best = candidates[0]
        if best:
            print(f"  ✅ danji_pages 캐시 히트: {best['complex_name']} ({best.get('location','')})")
        return best
    except Exception as e:
        print(f"  ⚠️ danji_pages 조회 실패: {e}")
        return None


def _danji_page_to_sales(danji_data):
    """danji_pages의 price_history를 hwik_engine 매매 데이터 형식으로 변환"""
    sales = []
    ph = danji_data.get("price_history") or {}
    for key, trades in ph.items():
        if "_" in key:  # 전세/월세 키는 스킵
            continue
        if not isinstance(trades, list):
            continue
        for t in trades:
            if not t.get("date") or not t.get("price"):
                continue
            sales.append({
                "dealAmount": str(t["price"]),
                "excluUseAr": key,
                "floor": str(t.get("floor", "")),
                "dealYear": t["date"][:4],
                "dealMonth": t["date"][5:7],
                "dealDay": t["date"][8:10] if len(t["date"]) >= 10 else "1",
                "aptNm": danji_data.get("complex_name", ""),
                "_parsed_price": t["price"],
                "_date": t["date"],
            })
    sales.sort(key=lambda x: x.get("_date", ""), reverse=True)
    return sales


def _danji_page_to_rent(danji_data):
    """danji_pages의 price_history에서 전세/월세 데이터 추출"""
    jeonse, wolse = [], []
    ph = danji_data.get("price_history") or {}
    for key, trades in ph.items():
        if not isinstance(trades, list):
            continue
        for t in trades:
            if not t.get("date") or not t.get("price"):
                continue
            row = {
                "deposit": str(t["price"]),
                "excluUseAr": key.replace("_jeonse", "").replace("_wolse", ""),
                "floor": str(t.get("floor", "")),
                "dealYear": t["date"][:4],
                "dealMonth": t["date"][5:7],
                "dealDay": t["date"][8:10] if len(t["date"]) >= 10 else "1",
                "aptNm": danji_data.get("complex_name", ""),
                "_parsed_price": t["price"],
                "_date": t["date"],
            }
            if "_wolse" in key:
                row["monthlyRent"] = str(t.get("monthly", 0))
                wolse.append(row)
            elif "_jeonse" in key:
                row["monthlyRent"] = "0"
                jeonse.append(row)
    jeonse.sort(key=lambda x: x.get("_date", ""), reverse=True)
    wolse.sort(key=lambda x: x.get("_date", ""), reverse=True)
    return jeonse, wolse


def _danji_page_to_nearby(danji_data):
    """danji_pages의 nearby_complex를 hwik_engine 형식으로 변환"""
    nearby_sales = {}
    nc = danji_data.get("nearby_complex") or []
    for n in nc:
        prices = n.get("prices") or {}
        # 가장 큰 면적의 가격을 대표로
        best_price = 0
        for k, v in prices.items():
            if v.get("price", 0) > best_price:
                best_price = v["price"]
        if best_price > 0:
            nearby_sales[n["name"]] = {
                "price": best_price,
                "dist": n.get("distance", 0),
                "lat": 0, "lon": 0,
                "trades": [{"dealAmount": str(best_price), "aptNm": n["name"]}],
            }
    return nearby_sales


SALES_API_URL      = 'http://apis.data.go.kr/1613000/RTMSDataSvcAptTradeDev/getRTMSDataSvcAptTradeDev'
RENT_API_URL       = 'http://apis.data.go.kr/1613000/RTMSDataSvcAptRent/getRTMSDataSvcAptRent'
OFFI_SALES_API_URL = 'https://apis.data.go.kr/1613000/RTMSDataSvcOffiTrade/getRTMSDataSvcOffiTrade'

# 단지명 앞 동 이름 접두어 제거용 (전국 주요 동/읍/리)
# RTMS API는 동 접두어 없이 저장하는 경우가 많아 키워드 매칭 보조로 사용
SHORT_DONGS = [
    # 서울 중랑구
    "신내","망우","묵동","면목","중화","상봉",
    # 서울 양천/강서/마포/은평
    "신정","목동","상암","수색","가좌","연희","홍제","홍은","불광","녹번","역촌","응암","대조","구산",
    # 서울 기타
    "잠실","둔촌","고덕","암사","명일","길동","천호","성내","풍납","가락","거여","마천",
    "공덕","아현","도화","용강","합정","망원","성산","상암","염리","대흥",
    "노량진","상도","대방","신대방","본동","흑석","동작",
    "방배","서초","반포","잠원","양재","원지",
    "개포","대치","역삼","도곡","일원","수서","자곡","율현","세곡",
    "압구정","청담","삼성","봉은사","논현","신사","학동","도산","신현",
    "등촌","화곡","방화","개화","공항","내발산","외발산","마곡",
    "봉천","신림","신원","서림","대학","난곡","난향","조원","미성",
    "사당","방배","낙성대",
    # 경기 주요 동/읍
    "갈매","별내","다산","강변","인창","교문","수택","토평","사노",  # 구리
    "덕소","와부","조안","진건","오남","퇴계원","화도","수동",        # 남양주
    "분당","수내","정자","서현","이매","야탑","판교","삼평","백현",   # 성남
    "일산","마두","백석","장항","주엽","대화","풍산","탄현","성석",   # 고양
    "동탄","반월","기흥","수지","구성","풍덕천","상현","광교",        # 화성/용인/수원
    "평촌","관양","비산","호계","범계","귀인",                        # 안양
    "중동","상동","부천","원미","심곡","도당","오정","소사",          # 부천
    # 인천
    "송도","구월","만수","간석","남동","논현","서창","연수","청학",
    # 기타 광역시
    "해운대","좌동","중동","우동","반여","반송","재송",               # 부산
    "수성","범어","만촌","황금","시지","고산",                        # 대구
    "죽전","구암","용현","학익","도화",                               # 인천/수원
]
OFFI_RENT_API_URL  = 'https://apis.data.go.kr/1613000/RTMSDataSvcOffiRent/getRTMSDataSvcOffiRent'


# ========================================================
# SSL 우회 — 정부 API(apis.data.go.kr)만 적용
# ========================================================
class TLSAdapter(HTTPAdapter):
    def init_poolmanager(self, *args, **kwargs):
        ctx = ssl.SSLContext(ssl.PROTOCOL_TLS_CLIENT)
        ctx.set_ciphers("DEFAULT@SECLEVEL=1")
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        kwargs['ssl_context'] = ctx
        return super().init_poolmanager(*args, **kwargs)

# 정부 API 전용 세션 (SSL 우회)
gov_session = requests.Session()
gov_session.mount("https://", TLSAdapter())
gov_session.verify = False

# 일반 세션 (카카오, Anthropic 등 — 정상 SSL)
session = requests.Session()


# ========================================================
# 유틸
# ========================================================
def _match_apt_name(kapt_name, trade_keys):
    """
    kaptName(공식명) → aptNm(API명) 매칭
    - trade 이름 전체가 kaptName에 포함되면 매칭 (API명이 짧은 경우 대응)
    - 예) '새한' in '신내새한아파트' → ✅
    - 예) '동성7' in '신내동성7차' → ✅
    - 예) '우남푸르' in '신내우남푸르미아' → ✅
    반환: (matched_key, score) — score는 매칭 글자수
    """
    nc = kapt_name.replace(" ", "")
    # prefix 제거 후보: 원본 + 앞 2~4글자 제거
    candidates = [nc] + [nc[skip:] for skip in range(2, 5) if len(nc) > skip + 1]

    best, best_score = None, 0
    for trade in trade_keys:
        if not trade or len(trade) < 2:
            continue
        score = 0
        for cand in candidates:
            # trade 전체가 cand에 포함 (API명이 짧은 경우)
            if trade in cand and len(trade) >= 2:
                s = len(trade)
                if s > score:
                    score = s
            # cand 앞부분이 trade에 포함
            for length in range(min(len(cand), len(trade), 7), 3, -1):
                if cand[:length] in trade:
                    if length > score:
                        score = length
                    break
        if score > best_score:
            best_score = score
            best = trade
    return best, best_score


def haversine(lat1, lon1, lat2, lon2):
    R = 6371000
    d_lat = math.radians(float(lat2) - float(lat1))
    d_lon = math.radians(float(lon2) - float(lon1))
    a = math.sin(d_lat/2)**2 + math.cos(math.radians(float(lat1))) * \
        math.cos(math.radians(float(lat2))) * math.sin(d_lon/2)**2
    return R * 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))


def _extract_keywords(name_clean):
    """단지명에서 검색 키워드 추출 (중복 로직 통합)"""
    GENERIC = {"아파트", "오피스텔", "빌라", "주상복합", "아파", "단지"}
    keywords = []
    kw1 = name_clean[:4]
    if kw1 not in GENERIC:
        keywords.append(kw1)
    if len(name_clean) > 4:
        kw2 = name_clean[2:6]
        if kw2 not in GENERIC and kw2 not in keywords:
            keywords.append(kw2)
    return keywords


# ── 한글 자모 분리 + 퍼지 매칭 ─────────────────────────
from difflib import SequenceMatcher

# 한글 유니코드 자모 테이블
_CHO  = "ㄱㄲㄴㄷㄸㄹㅁㅂㅃㅅㅆㅇㅈㅉㅊㅋㅌㅍㅎ"
_JUNG = "ㅏㅐㅑㅒㅓㅔㅕㅖㅗㅘㅙㅚㅛㅜㅝㅞㅟㅠㅡㅢㅣ"
_JONG  = [""] + list("ㄱㄲㄳㄴㄵㄶㄷㄹㄺㄻㄼㄽㄾㄿㅀㅁㅂㅄㅅㅆㅇㅈㅊㅋㅌㅍㅎ")

def _decompose_korean(text):
    """
    한글을 자모 단위로 분리
    '써' → 'ㅆㅓ', '서' → 'ㅅㅓ'
    영문/숫자/특수문자는 그대로 유지
    """
    result = []
    for ch in text:
        code = ord(ch) - 0xAC00
        if 0 <= code < 11172:
            cho  = code // 588
            jung = (code % 588) // 28
            jong = code % 28
            result.append(_CHO[cho])
            result.append(_JUNG[jung])
            if jong > 0:
                result.append(_JONG[jong])
        else:
            result.append(ch)
    return "".join(result)


def _fuzzy_score(name_a, name_b):
    """
    두 단지명의 유사도 점수 (0.0 ~ 1.0)
    한글 자모 분리 후 비교 — '써'와 '서'는 자모 1개 차이로 높은 점수
    """
    a = _decompose_korean(name_a.replace(" ", ""))
    b = _decompose_korean(name_b.replace(" ", ""))
    return SequenceMatcher(None, a, b).ratio()


def _fuzzy_match_best(search_name, candidates, min_score=0.82):
    """
    candidates 리스트에서 search_name과 가장 유사한 항목 반환
    candidates: [{"name": "xxx", ...}, ...] 또는 ["xxx", ...]
    반환: (best_item, score) 또는 (None, 0)
    """
    best, best_score = None, 0
    search_clean = search_name.replace(" ", "")
    for cand in candidates:
        cand_name = cand if isinstance(cand, str) else cand.get("kaptName", "")
        cand_clean = cand_name.replace(" ", "")

        # 길이 차이 너무 크면 건너뜀 (완전 다른 단지)
        if abs(len(search_clean) - len(cand_clean)) > 5:
            continue

        # 1순위: 정확 포함 — 짧은 입력(4자 이하)은 포함 매칭 금지
        if len(search_clean) >= 5 and (
            search_clean in cand_clean or cand_clean in search_clean
        ):
            score = 0.95 + (0.05 * (len(search_clean) / max(len(cand_clean), 1)))
        else:
            # 2순위: 퍼지 매칭
            score = _fuzzy_score(search_clean, cand_clean)

        if score > best_score:
            best_score = score
            best = cand
    if best_score >= min_score:
        return best, best_score
    return None, 0


def _extract_number_tag(name):
    """
    단지명에서 숫자 구분자 추출
    '한일써너스빌리젠시1단지' → '1단지'
    '우남퍼스트빌3차' → '3차'
    '래미안12블록' → '12블록'
    '효성써너스빌에코' → None (구분자 없음)
    """
    name_clean = name.replace(" ", "")
    m = re.search(r'(\d+)\s*(단지|차|동|블록|타워|단계|지구|차아파트|차오피스텔)', name_clean)
    if m:
        return m.group(0).replace(" ", "")
    # "1단지" 패턴이 아니라 끝에 숫자만 있는 경우 ("동성7" 등)
    m2 = re.search(r'(\d+)\s*$', name_clean)
    if m2:
        return m2.group(0)
    return None


def _extract_road_info(doro_juso):
    """
    도로명주소에서 도로명 + 본번 추출
    '서울특별시 중랑구 망우로60길 37' → ('망우로60길', '37')
    '망우로60길 37' → ('망우로60길', '37')
    '' → (None, None)
    """
    if not doro_juso or not doro_juso.strip():
        return None, None
    parts = doro_juso.strip().split()
    road_nm = None
    bonbun = None
    for i, p in enumerate(parts):
        # '로' 또는 '길'이 포함된 파트가 도로명
        if re.search(r'[로길]', p) and re.search(r'\d', p):
            road_nm = p
            # 다음 파트가 본번
            if i + 1 < len(parts):
                bm = re.match(r'(\d+)', parts[i + 1])
                if bm:
                    bonbun = bm.group(1)
            break
    return road_nm, bonbun


def _match_row(row, keywords, number_tag, road_nm, road_bonbun, name_field,
               search_name=None, fuzzy_threshold=0.82):
    """
    실거래가 API 응답 1건이 검색 대상 단지와 일치하는지 판별
    1단계: 키워드 AND 매칭 (기존) — 빠름
    1-B단계: 키워드 실패 시 퍼지 매칭 (search_name 필요) — 오타 허용
    2단계: 숫자 구분자 필터 (있을 때만 + 역방향 체크)
    3단계: 도로명주소 대조 (있을 때만)
    """
    nm = row.get(name_field, "").replace(" ", "")

    # 1단계: 키워드 AND (빠른 필터)
    keyword_ok = keywords and all(kw in nm for kw in keywords)

    # 1-B단계: 키워드 실패 → 퍼지 매칭 시도 (짧은 이름은 퍼지 금지)
    if not keyword_ok:
        if search_name and len(search_name) >= 6:
            # 길이 차이 너무 크면 건너뜀
            if abs(len(search_name) - len(nm)) > 5:
                return False
            score = _fuzzy_score(search_name, nm)
            if score >= fuzzy_threshold:
                keyword_ok = True
        if not keyword_ok:
            return False

    # 2단계: 숫자 구분자
    if number_tag:
        # 입력에 "1단지"가 있는데 응답에 없으면 탈락
        if number_tag not in nm:
            return False
    else:
        # 입력에 숫자가 없는데 응답에 숫자 구분자가 있으면 탈락
        # 예: 입력 "래미안블레스티지" → "래미안블레스티지1단지" 매칭 방지
        if search_name and not re.search(r'\d', search_name):
            nm_number = _extract_number_tag(nm)
            if nm_number:
                return False

    # 3단계: 도로명주소 — API 응답에 roadNm 필드가 있으면 대조
    if road_nm:
        row_road = row.get("roadNm", "").replace(" ", "")
        if row_road:
            # 도로명이 서로 포함 관계여야 매칭
            if road_nm not in row_road and row_road not in road_nm:
                return False
            # 본번도 있으면 추가 대조 (앞자리 0 제거 후 비교: '00085' == '85')
            if road_bonbun:
                row_bonbun = str(row.get("roadNmBonbun", "")).strip().lstrip("0") or "0"
                cmp_bonbun = road_bonbun.lstrip("0") or "0"
                if row_bonbun and row_bonbun != cmp_bonbun:
                    return False

    return True


def _match_by_road(row, road_nm, road_bonbun):
    """
    도로명 + 본번으로 매칭 (매매 PRIMARY)
    roadNm, roadNmBonbun 필드 사용
    """
    if not road_nm:
        return False
    row_road = row.get("roadNm", "").replace(" ", "")
    if not row_road:
        return False
    # 도로명 포함 관계 확인
    if road_nm not in row_road and row_road not in road_nm:
        return False
    # 본번 확인 (앞자리 0 제거 후 비교)
    if road_bonbun:
        row_bonbun = str(row.get("roadNmBonbun", "")).strip().lstrip("0") or "0"
        cmp_bonbun = road_bonbun.lstrip("0") or "0"
        if row_bonbun and row_bonbun != cmp_bonbun:
            return False
    return True


def _match_by_jibun(row, jibun, umd_nm=""):
    """
    지번으로 매칭 (전월세 PRIMARY — 전월세 API에 roadNm 없음)
    jibun 필드 사용
    """
    if not jibun:
        return False
    row_jibun = row.get("jibun", "").strip()
    if not row_jibun:
        return False
    # 지번 정규화: "591-0" → "591", "00816" → "816"
    def _norm(j):
        parts = j.split("-")
        main = parts[0].lstrip("0") or "0"
        sub  = parts[1].lstrip("0") if len(parts) > 1 else "0"
        return (main, sub)
    if _norm(row_jibun) != _norm(jibun):
        return False
    # 동 이름도 있으면 추가 확인
    if umd_nm:
        row_umd = row.get("umdNm", "").replace(" ", "")
        if row_umd and umd_nm not in row_umd and row_umd not in umd_nm:
            return False
    return True


def parse_price(val):
    """'12,000' → 12000"""
    try:
        return int(str(val).replace(',', '').strip())
    except (ValueError, TypeError):
        return None


def format_price_word(val):
    """12000 → '1억 2,000만원'"""
    n = parse_price(val)
    if n is None:
        return "정보없음"
    eok  = n // 10000
    man  = n % 10000
    if eok > 0 and man > 0:
        return f"{eok}억 {man:,}만원"
    elif eok > 0:
        return f"{eok}억원"
    else:
        return f"{man:,}만원"

def normalize_row(row, source_type):
    """
    API별 필드명 통일
    source_type: "apt" | "offi"
    offiNm → aptNm 로 변환, 나머지는 필드명 동일하므로 그대로
    """
    r = dict(row)
    if source_type == "offi":
        if "offiNm" in r:
            r["aptNm"] = r.pop("offiNm")
    return r

def _is_valid_num(val):
    """유효한 숫자 문자열인지 확인 (preDeposit/preMonthlyRent 판별용)"""
    try:
        return bool(val) and float(str(val).replace(',', '').strip()) > 0
    except (ValueError, TypeError):
        return False


def _resolve_deposit(row):
    """전세 보증금: deposit 우선, 없으면 preDeposit 사용"""
    dep = row.get("deposit", "")
    pre = row.get("preDeposit", "")
    if _is_valid_num(dep):
        return dep, "deposit"
    if _is_valid_num(pre):
        return pre, "preDeposit"
    return "", "none"


def _resolve_monthly(row):
    """월세: (보증금, 월세) 튜플 반환"""
    dep, _ = _resolve_deposit(row)
    mon = row.get("monthlyRent", "")
    pre_mon = row.get("preMonthlyRent", "")
    if _is_valid_num(mon):
        return dep, mon
    if _is_valid_num(pre_mon):
        return dep, pre_mon
    return dep, ""


def get_pyeong(exclu_ar, build_year=None):
    """전용면적 → 공급면적 기준 평형 (비율 추정 — 건축물대장 데이터 없을 때 fallback)"""
    try:
        ar = float(str(exclu_ar).replace(',', '').strip())
        if build_year:
            yr = int(str(build_year)[:4])
            ratio = 0.73 if yr < 2000 else (0.80 if yr < 2010 else 0.84)
        else:
            ratio = 0.80
        supply = ar / ratio
        return round(supply / 3.305785)
    except (ValueError, TypeError) as e:
        print(f"  ⚠️ 평형 변환 오류 (exclu_ar={exclu_ar}): {e}")


# ========================================================
# 건축물대장 공급면적 조회 (getBrExposPubuseAreaInfo)
# ========================================================
BUILDING_API_URL = "https://apis.data.go.kr/1613000/BldRgstHubService/getBrExposPubuseAreaInfo"
APT_TYPES_SUPPLY  = {"아파트", "공동주택", "연립주택", "다세대주택"}
OFFI_TYPES_SUPPLY = {"오피스텔"}


def _cluster_supply_areas(area_map, tolerance=2.0):
    """
    전용면적 tolerance㎡ 이내 클러스터링
    {59.49: [121.3], 59.60: [121.4], 59.71: [121.5]} → {59.60: (121.4, 14)}
    """
    from collections import defaultdict
    keys = sorted(area_map.keys())
    if not keys:
        return {}
    clusters, current = [], [keys[0]]
    for k in keys[1:]:
        if k - current[0] <= tolerance:
            current.append(k)
        else:
            clusters.append(current)
            current = [k]
    clusters.append(current)
    result = {}
    for cluster in clusters:
        rep   = max(cluster, key=lambda k: len(area_map[k]))
        all_s = []
        for k in cluster:
            all_s.extend(area_map[k])
        result[round(rep, 2)] = (round(sum(all_s) / len(all_s), 2), len(all_s))
    return result


def _calc_raw_supply_map(items):
    """호별 전용+공용 합산 → {전용면적: [공급면적, ...]}"""
    from collections import defaultdict
    ho_expos, ho_pubuse = {}, {}
    for it in items:
        gb = it.get("exposPubuseGbCd", "")
        ho = it.get("hoNm", "").strip()
        try:
            ar = float(it.get("area", 0) or 0)
        except:
            ar = 0
        if not ho or ar <= 0:
            continue
        if gb == "1":
            ho_expos[ho]  = ho_expos.get(ho, 0) + ar
        elif gb == "2":
            ho_pubuse[ho] = ho_pubuse.get(ho, 0) + ar
    raw = defaultdict(list)
    for ho, expos in ho_expos.items():
        pubuse = ho_pubuse.get(ho, 0)
        if pubuse > 0:
            raw[round(expos, 2)].append(round(expos + pubuse, 2))
    return raw


def _get_supply_map_for_nearby(apt_name):
    """
    주변 단지명 → Supabase DB의 pyeongs 컬럼에서 공급면적 맵 반환
    pyeongs 없으면 doro_juso로 건축물대장 API 폴백
    캐시 적용 (같은 단지 중복 호출 방지)
    """
    if not hasattr(_get_supply_map_for_nearby, "_cache"):
        _get_supply_map_for_nearby._cache = {}
    if apt_name in _get_supply_map_for_nearby._cache:
        return _get_supply_map_for_nearby._cache[apt_name]

    result = {}
    try:
        name_clean = apt_name.replace(" ", "")
        res = requests.get(
            f"{SUPABASE_URL}/rest/v1/apartments",
            headers=SUPABASE_HEADERS,
            params={"select": "doro_juso,property_type,pyeongs", "slug": f"ilike.*{name_clean}*", "limit": "1"},
            timeout=5,
        )
        rows = res.json() if res.status_code == 200 else []
        if isinstance(rows, list) and rows:
            pyeongs = rows[0].get("pyeongs")
            if pyeongs and isinstance(pyeongs, list) and len(pyeongs) > 0:
                # DB에 평형정보 있으면 바로 사용 (API 호출 불필요)
                result = {round(float(p["exclu"]), 2): round(float(p["supply"]), 2)
                          for p in pyeongs if p.get("exclu") and p.get("supply")}
            else:
                # pyeongs 없으면 건축물대장 API 폴백
                doro  = rows[0].get("doro_juso", "")
                ptype = rows[0].get("property_type", "apt")
                if doro:
                    result = get_supply_area_map(doro, property_type=ptype)
    except Exception as e:
        print(f"  ⚠️ 주변단지 공급면적 조회 오류 ({apt_name}): {e}")

    _get_supply_map_for_nearby._cache[apt_name] = result
    return result


def get_supply_area_map(road_address, property_type="apt"):
    """
    도로명 주소 → {전용면적: 공급면적} 딕셔너리 반환
    property_type: "apt" | "offi"
    반환: {18.5: 31.63, 59.61: 81.05, ...}  — 없으면 {}
    """
    if not road_address or not road_address.strip():
        return {}

    print(f"\n공급면적 조회 중: {road_address}")

    # 카카오 주소 검색으로 b_code / 지번 추출
    try:
        url = "https://dapi.kakao.com/v2/local/search/address.json"
        headers = {"Authorization": f"KakaoAK {KAKAO_API_KEY}"}
        res = session.get(url, headers=headers, params={"query": road_address}, timeout=10)
        docs = res.json().get("documents", [])
        if not docs:
            print("  ⚠️ 공급면적: 주소 검색 실패 → 건너뜀")
            return {}
        ai     = docs[0].get("address") or {}
        b_code = ai.get("b_code", "")
        main   = ai.get("main_address_no", "")
        sub    = ai.get("sub_address_no", "0")
        if not b_code:
            print("  ⚠️ 공급면적: b_code 없음 → 건너뜀")
            return {}
        sigunguCd = b_code[:5]
        bjdongCd  = b_code[5:]
        bun       = main.zfill(4)
        ji        = sub.zfill(4) if sub and sub != "0" else "0000"
    except Exception as e:
        print(f"  ⚠️ 공급면적 주소 변환 오류: {e}")
        return {}

    # 건축물대장 API 호출 (최대 10페이지 = 1000건, 평형 파악 목적)
    base_params = {
        "serviceKey": GOV_SERVICE_KEY,
        "sigunguCd":  sigunguCd,
        "bjdongCd":   bjdongCd,
        "bun":        bun,
        "ji":         ji,
        "numOfRows":  "100",
        "_type":      "json",
    }

    def _parse(body):
        items = body.get("items") or {}
        lst   = items.get("item", [])
        return [lst] if isinstance(lst, dict) else (lst or [])

    try:
        r     = gov_session.get(BUILDING_API_URL,
                                params={**base_params, "pageNo": "1"}, timeout=15)
        body  = r.json().get("response", {}).get("body", {})
        total = int(body.get("totalCount", 0))
        if total == 0:
            print("  ⚠️ 공급면적: 건축물대장 레코드 없음")
            return {}

        item_list   = _parse(body)
        total_pages = math.ceil(total / 100)
        max_pages   = min(total_pages, 10)

        for page in range(2, max_pages + 1):
            r = gov_session.get(BUILDING_API_URL,
                                params={**base_params, "pageNo": str(page)}, timeout=15)
            item_list.extend(_parse(r.json().get("response", {}).get("body", {})))

        # property_type에 맞는 항목만 필터
        type_set = OFFI_TYPES_SUPPLY if property_type == "offi" else APT_TYPES_SUPPLY
        filtered = [it for it in item_list if it.get("mainPurpsCdNm", "") in type_set]

        if not filtered:
            print(f"  ⚠️ 공급면적: '{property_type}' 유형 레코드 없음 (전체 {len(item_list)}건)")
            return {}

        raw     = _calc_raw_supply_map(filtered)
        cluster = _cluster_supply_areas(raw, tolerance=2.0)

        # {전용면적: 공급면적} 형태로 변환 (호수 정보 제거)
        result = {expos: avg_s for expos, (avg_s, _) in cluster.items()}
        print(f"  ✅ 공급면적 {len(result)}개 평형 확인")
        for expos, supply in sorted(result.items()):
            print(f"     전용 {expos:.2f}㎡ → 공급 {supply:.2f}㎡")
        return result

    except Exception as e:
        print(f"  ⚠️ 공급면적 조회 오류: {e}")
        return {}


def get_supply_from_map(exclu_ar, supply_map):
    """
    전용면적 → 공급면적 변환
    supply_map에서 가장 가까운 전용면적 매핑값 반환 (2㎡ 이내)
    없으면 None
    """
    if not supply_map or exclu_ar is None:
        return None
    try:
        ar = float(str(exclu_ar).replace(',', '').strip())
    except:
        return None
    best_key  = min(supply_map.keys(), key=lambda k: abs(k - ar))
    if abs(best_key - ar) <= 2.0:
        return supply_map[best_key]
    return None


# ========================================================
# 카카오 / 주소 API
# ========================================================
def address_to_info(address):
    """주소 → (lat, lon, b_code, jibun, umd_nm)"""
    url = "https://dapi.kakao.com/v2/local/search/address.json"
    headers = {"Authorization": f"KakaoAK {KAKAO_API_KEY}"}
    try:
        res = session.get(url, headers=headers, params={"query": address}, timeout=10)
        docs = res.json().get("documents", [])
        if docs:
            doc = docs[0]
            addr_info = doc.get("address") or {}
            b_code  = addr_info.get("b_code", "")
            lat     = float(doc.get("y", 0))
            lon     = float(doc.get("x", 0))
            main_no = addr_info.get("main_address_no", "")
            sub_no  = addr_info.get("sub_address_no", "")
            jibun   = f"{main_no}-{sub_no}" if sub_no and sub_no != "0" else main_no
            umd_nm  = addr_info.get("region_3depth_name", "")
            return lat, lon, b_code, jibun, umd_nm
    except Exception as e:
        print(f"❌ 주소변환 오류: {e}")
    return None, None, None, None, None


def search_kakao_keyword(query, allow_offi=False):
    url = "https://dapi.kakao.com/v2/local/search/keyword.json"
    headers = {"Authorization": f"KakaoAK {KAKAO_API_KEY}"}
    try:
        res = session.get(url, headers=headers, params={"query": query, "size": 5}, timeout=10)
        docs = res.json().get("documents", [])
        if allow_offi:
            # 오피스텔 검색: 오피스텔 카테고리 우선, 없으면 아파트, 없으면 첫번째
            offi = [d for d in docs if "오피스텔" in d.get("category_name", "")]
            apts = [d for d in docs if "아파트" in d.get("category_name", "")]
            return offi[0] if offi else (apts[0] if apts else (docs[0] if docs else None))
        else:
            # 아파트 검색: 아파트 카테고리 우선, 없으면 첫번째
            apts = [d for d in docs if "아파트" in d.get("category_name", "")]
            return apts[0] if apts else (docs[0] if docs else None)
    except Exception as e:
        print(f"❌ 카카오 검색 오류: {e}")
    return None


# ========================================================
# 학교 정보 (카카오 SC4 카테고리)
# ========================================================
# ========================================================
# 학교 정보 (Supabase schools 테이블 — 카카오 API 대체)
# ========================================================
def get_nearby_schools(lat, lon, radius=500):
    """
    Supabase schools 테이블에서 반경 내 학교 조회
    haversine 거리 계산 → radius(m) 이내 필터
    반환: [{"name": ..., "type": ..., "dist": ..., "address": ...}, ...]
    """
    if not lat or not lon:
        return []

    # 위도 1도 ≈ 111km, 경도 1도 ≈ 88km(서울 기준)
    # 넉넉하게 bbox로 1차 필터
    margin = radius / 90000.0
    try:
        # Supabase REST: 복수 필터는 params list로 처리
        res = requests.get(
            f"{SUPABASE_URL}/rest/v1/schools",
            headers=SUPABASE_HEADERS,
            params=[
                ("select",  "name,type,address,lat,lon"),
                ("lat",     f"gte.{lat - margin}"),
                ("lat",     f"lte.{lat + margin}"),
                ("lon",     f"gte.{lon - margin}"),
                ("lon",     f"lte.{lon + margin}"),
                ("limit",   "50"),
            ],
            timeout=10,
        )
        candidates = res.json() if res.status_code == 200 else []
        if not isinstance(candidates, list):
            candidates = []
    except Exception as e:
        print(f"❌ 학교 DB 조회 오류: {e}")
        return []

    # 2차 필터: 정확한 haversine 거리 계산
    schools = []
    for s in candidates:
        try:
            dist = haversine(lat, lon, float(s["lat"]), float(s["lon"]))
        except (ValueError, TypeError, KeyError):
            continue
        if dist <= radius:
            schools.append({
                "name":    s.get("name", ""),
                "type":    s.get("type", "기타"),
                "dist":    dist,
                "address": s.get("address", ""),
            })

    schools.sort(key=lambda x: x["dist"])
    print(f"✅ 주변 학교 {len(schools)}개 발견 (반경 {radius}m, DB 조회)")
    return schools


# ========================================================
# 지하철/철도역 정보 (Supabase stations 테이블)
# ========================================================
def get_nearby_stations(lat, lon, radius=1000):
    """
    Supabase stations 테이블에서 반경 내 역 조회
    도보 속도: 80m/분 (빠른 걸음 기준 — 실제보다 약간 빠르게 표시)
    반환: [{"name": ..., "line": ..., "dist": ..., "walk_min": ...}, ...]
    """
    if not lat or not lon:
        return []

    margin = radius / 90000.0
    try:
        res = requests.get(
            f"{SUPABASE_URL}/rest/v1/stations",
            headers=SUPABASE_HEADERS,
            params=[
                ("select", "name,line,operator,lat,lon"),
                ("lat",    f"gte.{lat - margin}"),
                ("lat",    f"lte.{lat + margin}"),
                ("lon",    f"gte.{lon - margin}"),
                ("lon",    f"lte.{lon + margin}"),
                ("limit",  "30"),
            ],
            timeout=10,
        )
        candidates = res.json() if res.status_code == 200 else []
        if not isinstance(candidates, list):
            candidates = []
    except Exception as e:
        print(f"❌ 역 DB 조회 오류: {e}")
        return []

    stations = []
    for s in candidates:
        try:
            dist = haversine(lat, lon, float(s["lat"]), float(s["lon"]))
        except:
            continue
        if dist <= radius:
            walk_min = max(1, round(dist / 80))
            stations.append({
                "name":     s.get("name", ""),
                "line":     s.get("line", ""),
                "dist":     dist,
                "walk_min": walk_min,
            })

    stations.sort(key=lambda x: x["dist"])
    print(f"✅ 주변 역 {len(stations)}개 발견 (반경 {radius}m)")
    return stations


# ========================================================
# 단지 목록 / 상세
# ========================================================
def get_apt_list(bjd_code):
    url = "https://apis.data.go.kr/1613000/AptListService3/getLegaldongAptList3"
    for code in [bjd_code, bjd_code[:8], bjd_code[:5]]:
        try:
            res = gov_session.get(url, params={
                "serviceKey": GOV_SERVICE_KEY,
                "bjdCode": code, "numOfRows": 9999, "pageNo": 1
            }, timeout=15)
            body = res.json().get("response", {}).get("body", {})
            items = body.get("items")
            if not items:
                continue
            if isinstance(items, dict):
                items = items.get("item", [])
            if isinstance(items, dict):
                items = [items]
            result = [{"kaptCode": i.get("kaptCode", ""), "kaptName": i.get("kaptName", "")}
                      for i in (items or []) if i.get("kaptCode")]
            if result:
                print(f"✅ {len(result)}개 단지 조회됨 (bjdCode: {code})")
                return result
        except Exception as e:
            print(f"  오류({code}): {e}")
    return []


def get_apt_detail(kapt_code):
    url = "https://apis.data.go.kr/1613000/AptBasisInfoServiceV4/getAphusBassInfoV4"
    try:
        res = gov_session.get(url, params={"serviceKey": GOV_SERVICE_KEY, "kaptCode": kapt_code},
                         timeout=15)
        item = res.json().get("response", {}).get("body", {}).get("item", {})
        if item:
            return {
                "kaptCode":       item.get("kaptCode", ""),
                "kaptName":       item.get("kaptName", ""),
                "doroJuso":       item.get("doroJuso", ""),
                "kaptdaCnt":      item.get("kaptdaCnt", ""),
                "kaptUsedate":    item.get("kaptUsedate", ""),
                "kaptTopFloor":   item.get("kaptTopFloor", ""),
                "kaptBcompany":   item.get("kaptBcompany", ""),
                "kaptTotPkCnt":   item.get("kaptTotPkCnt", ""),
                "hhldCount":      item.get("hhldCount", ""),
            }
    except Exception as e:
        print(f"❌ 상세 조회 오류: {e}")
    return None


# ========================================================
# Claude 파싱
# ========================================================
def parse_input_with_claude(user_input):
    print("Claude로 입력 파싱 중...")
    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers=anthropic_headers(),
            json={
                "model": ANTHROPIC_MODEL_HAIKU,
                "max_tokens": 300,
                "messages": [{
                    "role": "user",
                    "content": f"""아래 입력에서 부동산 정보를 추출해서 JSON으로만 반환하세요.
반드시 JSON만 반환하고 다른 텍스트는 없어야 합니다.

중요: 단지명에 오타가 있을 수 있습니다. 가장 가능성 높은 정식 명칭으로 교정하세요.
- 쌍자음/단자음 혼동: 써→서, 빌→빌 등 → 일반적인 아파트/오피스텔 브랜드명 기준으로 판단
- 영문/한글 혼용: e편한세상→이편한세상, SK뷰→SK VIEW 등 → 가장 널리 쓰이는 표기로 통일
- 띄어쓰기/축약: 래미안블레스티지→래미안 블레스티지 등
- 단지 구분 생략: "한일써너스빌리젠시" → 1단지/2단지 구분이 불명확하면 그대로 두세요

property_hint: 입력에 "오피스텔"이 명시되어 있으면 "officetel", 아니면 "apt"

입력: {user_input}

JSON 형식:
{{
  "gu": "구이름(예:중랑구)",
  "dong": "동이름(예:신내동)",
  "apt_name": "단지명 — 오타 교정된 정식 명칭(예:영풍마드레빌)",
  "road_address": "도로명주소 전체(없으면 빈문자열)",
  "property_hint": "apt 또는 officetel"
}}"""
                }]
            },
            timeout=15
        )
        data = res.json()
        text = data.get("content", [{}])[0].get("text", "").strip()
        if "```" in text:
            text = text.split("```")[1].replace("json", "").strip()
        result = json.loads(text)
        print(f"✅ 파싱 결과: {result}")
        return result
    except Exception as e:
        print(f"❌ Claude 파싱 오류: {e}")
        return None


# ========================================================
# 실거래가 조회 (매매) — 아파트 → 오피스텔 폭포식
# ========================================================
def fetch_sales_auto(kapt_name, lawd_cd, months=36, force_offi=False, road_address="",
                     jibun="", umd_nm=""):
    """
    매매 실거래가 자동 탐지
    아파트: 1순위 도로명+본번, 2순위 jibun, 3순위 키워드
    오피스텔: 키워드 PRIMARY
    반환: (data_list, property_type)   property_type = "apt" | "offi" | None
    """
    now = datetime.now()
    name_clean = kapt_name.replace(" ", "")


    # 아파트/오피스텔 접미사 제거 — RTMS API는 단지명에 접미사 없이 저장
    # 예) '신내새한아파트' → '신내새한', '효성써너스빌에코오피스텔' → '효성써너스빌에코'
    for suffix in ('아파트', '오피스텔', '주상복합'):
        if name_clean.endswith(suffix) and len(name_clean) > len(suffix) + 2:
            name_clean = name_clean[:-len(suffix)]
            break
    number_tag = _extract_number_tag(kapt_name)
    road_nm, road_bonbun = _extract_road_info(road_address)

    # number_tag 제거 후 키워드 추출
    # 예) 신내우디안1단지 → number_tag='1단지' → name_no_tag='신내우디안' → ['신내우디','우디안']
    name_no_tag = name_clean.replace(number_tag, "") if number_tag else name_clean
    keywords = _extract_keywords(name_no_tag) or _extract_keywords(name_clean)
    if not keywords:
        print(f"  ⚠️ 유효 키워드 없음 (단지명이 너무 일반적): {kapt_name}")
        return [], None

    # 동 이름 제거 버전 키워드도 준비 (신내동성4차 → 동성4차)
    # SHORT_DONGS는 모듈 상수 사용
    name_short = name_no_tag  # number_tag 제거 버전 기준으로 동 접두어 제거
    for dong in SHORT_DONGS:
        if name_no_tag.startswith(dong) and len(name_no_tag) > len(dong) + 1:
            name_short = name_no_tag[len(dong):]
            break
    keywords_short = _extract_keywords(name_short) if name_short != name_no_tag else []

    if number_tag:
        print(f"  숫자 구분자: {number_tag}")
    if road_nm:
        print(f"  도로명 필터: {road_nm} {road_bonbun or ''}")

    # ── [DEBUG] 환경변수 HWIK_DEBUG=1 설정 시 파라미터 출력 ──
    _debug = os.environ.get("HWIK_DEBUG")
    if _debug:
        print(f"\n  ╔═ DEBUG fetch_sales_auto ══════════════════")
        print(f"  ║ 입력 단지명 : {kapt_name!r}")
        print(f"  ║ name_no_tag : {name_no_tag!r}")
        print(f"  ║ keywords    : {keywords}")
        print(f"  ║ keywords_short: {keywords_short}")
        print(f"  ║ number_tag  : {number_tag!r}")
        print(f"  ╚══════════════════════════════════════════")

    def _fetch(api_url, name_field, label):
        print(f"[매매] {label} API 조회 중 ({months}개월) / 키워드: {keywords}")
        ym_list = [(now - relativedelta(months=i)).strftime("%Y%m") for i in range(months)]

        def _fetch_one(ym):
            rows = []
            try:
                resp = gov_session.get(api_url, params={
                    "serviceKey": GOV_SERVICE_KEY,
                    "LAWD_CD": lawd_cd, "DEAL_YMD": ym,
                    "pageNo": "1", "numOfRows": "9999"
                }, headers=HDR_XML, timeout=30)
                if resp.status_code == 200 and resp.content:
                    root = ET.fromstring(resp.content)
                    _debug_seen = set()
                    for item in root.findall(".//item"):
                        row = {c.tag: (c.text.strip() if c.text else "") for c in item}
                        nm_raw = row.get(name_field, "")

                        # 1순위: 도로명+본번 (가장 정확, 구형단지는 roadNm 없을 수 있음)
                        matched = _match_by_road(row, road_nm, road_bonbun) if road_nm else False

                        # 2순위: 지번 + 키워드 2차 필터
                        # (같은 지번에 여러 단지가 있을 수 있으므로 키워드로 재확인)
                        if not matched and jibun:
                            if _match_by_jibun(row, jibun, umd_nm):
                                nm_clean = nm_raw.replace(" ", "")
                                # 키워드 중 하나라도 단지명에 포함되면 매칭
                                kw_ok = any(kw in nm_clean for kw in keywords + keywords_short)
                                # 키워드가 없으면 지번만으로 매칭 (키워드 추출 실패 대비)
                                matched = kw_ok or not (keywords + keywords_short)

                        # 3순위: 키워드 (도로명/지번 모두 없는 경우)
                        if not matched and not road_nm and not jibun:
                            matched = _match_row(row, keywords, number_tag, None, None,
                                          name_field, search_name=name_no_tag)
                            if not matched and keywords_short:
                                matched = _match_row(row, keywords_short, number_tag, None, None,
                                              name_field, search_name=name_short)

                        # [DEBUG]
                        if _debug:
                            if matched or any(kw in nm_raw.replace(" ", "")
                                              for kw in keywords + keywords_short):
                                _debug_seen.add(
                                    f"  {'✅' if matched else '❌'} [{ym}] {name_field}={nm_raw!r}"
                                    f"  road={row.get('roadNm','')} {row.get('roadNmBonbun','')} jibun={row.get('jibun','')}")
                        if matched:
                            rows.append(row)
                    if _debug:
                        for s in sorted(_debug_seen):
                            print(s)
            except Exception as e:
                print(f"  [오류] {ym}: {e}")
            return rows

        all_data = []
        with ThreadPoolExecutor(max_workers=4) as pool:
            futures = {pool.submit(_fetch_one, ym): ym for ym in ym_list}
            for f in as_completed(futures):
                all_data.extend(f.result())

        # [DEBUG] 매칭된 단지명 종류 요약
        if _debug and all_data:
            unique_names = Counter(r.get(name_field, "") for r in all_data)
            print(f"\n  ╔═ DEBUG 매칭 결과 요약 ({label}) ════════════")
            for nm, cnt in unique_names.most_common():
                print(f"  ║  {nm!r}: {cnt}건")
            print(f"  ╚══════════════════════════════════════════")

        return all_data

    # 1순위: 아파트 (force_offi=True면 건너뜀)
    if not force_offi:
        data = _fetch(SALES_API_URL, "aptNm", "아파트")
        if data:
            print(f"✅ 아파트 매매 {len(data)}건")
            return [normalize_row(r, "apt") for r in data], "apt"
        print("  → 아파트 0건. 오피스텔 API 시도...")
    else:
        print("  [오피스텔 직행] 아파트 API 생략")

    # 2순위: 오피스텔
    data = _fetch(OFFI_SALES_API_URL, "offiNm", "오피스텔")
    if data:
        print(f"✅ 오피스텔 매매 {len(data)}건")
        return [normalize_row(r, "offi") for r in data], "offi"

    print("  ⚠️ 매매 데이터 없음 (아파트/오피스텔 모두 0건)")
    return [], None


# ========================================================
# 실거래가 조회 (전월세) — property_type 기반, 폭포식 안전망
# ========================================================
def fetch_rent_auto(kapt_name, lawd_cd, property_type, months=36, road_address="",
                    jibun="", umd_nm=""):
    """
    전월세 실거래가
    아파트: jibun PRIMARY (전월세 API에 roadNm 없음)
    오피스텔: 키워드 PRIMARY
    반환: (jeonse_list, wolse_list)
    """
    now = datetime.now()
    name_clean = kapt_name.replace(" ", "")

    number_tag = _extract_number_tag(kapt_name)
    road_nm, road_bonbun = _extract_road_info(road_address)

    # number_tag 제거 후 키워드 추출
    name_no_tag = name_clean.replace(number_tag, "") if number_tag else name_clean
    keywords = _extract_keywords(name_no_tag) or _extract_keywords(name_clean)
    if not keywords:
        print(f"  ⚠️ 유효 키워드 없음: {kapt_name}")
        return [], []

    # 동 이름 제거 버전
    # SHORT_DONGS는 모듈 상수 사용
    name_short = name_no_tag
    for dong in SHORT_DONGS:
        if name_no_tag.startswith(dong) and len(name_no_tag) > len(dong) + 1:
            name_short = name_no_tag[len(dong):]
            break
    keywords_short = _extract_keywords(name_short) if name_short != name_no_tag else []

    def _fetch(api_url, name_field, label, use_jibun=False):
        print(f"[전월세] {label} API 조회 중 ({months}개월) / "
              f"{'지번: ' + jibun if use_jibun and jibun else '키워드: ' + str(keywords)}")
        ym_list = [(now - relativedelta(months=i)).strftime("%Y%m") for i in range(months)]

        def _fetch_one(ym):
            rows = []
            try:
                resp = gov_session.get(api_url, params={
                    "serviceKey": GOV_SERVICE_KEY,
                    "LAWD_CD": lawd_cd, "DEAL_YMD": ym,
                    "pageNo": "1", "numOfRows": "9999"
                }, headers=HDR_XML, timeout=30)
                if resp.status_code == 200 and resp.content:
                    root = ET.fromstring(resp.content)
                    for item in root.findall(".//item"):
                        row = {c.tag: (c.text.strip() if c.text else "") for c in item}

                        if use_jibun and jibun:
                            # ── PRIMARY: 지번 매칭 (아파트 전월세) ──
                            matched = _match_by_jibun(row, jibun, umd_nm)
                        else:
                            # ── PRIMARY: 키워드 매칭 (오피스텔 전월세) ──
                            matched = _match_row(row, keywords, number_tag, road_nm, road_bonbun,
                                          name_field, search_name=name_no_tag)
                            if not matched and keywords_short:
                                matched = _match_row(row, keywords_short, number_tag, road_nm, road_bonbun,
                                              name_field, search_name=name_short)

                        if matched:
                            rows.append(row)
            except Exception as e:
                print(f"  [오류] {ym}: {e}")
            return rows

        all_data = []
        with ThreadPoolExecutor(max_workers=4) as pool:
            futures = {pool.submit(_fetch_one, ym): ym for ym in ym_list}
            for f in as_completed(futures):
                all_data.extend(f.result())
        return all_data

    def _split(all_data, src_type):
        normalized = [normalize_row(r, src_type) for r in all_data]
        jeonse = [r for r in normalized if not r.get("monthlyRent") or r["monthlyRent"] == "0"]
        wolse  = [r for r in normalized if r.get("monthlyRent") and r["monthlyRent"] != "0"]
        return jeonse, wolse

    if property_type == "apt":
        data = _fetch(RENT_API_URL, "aptNm", "아파트", use_jibun=True)
        jeonse, wolse = _split(data, "apt")
        print(f"✅ 전세 {len(jeonse)}건 / 월세 {len(wolse)}건")
        return jeonse, wolse

    if property_type == "offi":
        data = _fetch(OFFI_RENT_API_URL, "offiNm", "오피스텔", use_jibun=False)
        jeonse, wolse = _split(data, "offi")
        print(f"✅ 오피스텔 전세 {len(jeonse)}건 / 월세 {len(wolse)}건")
        return jeonse, wolse

    # 안전망 — property_type 없는 경우
    data = _fetch(RENT_API_URL, "aptNm", "아파트")
    if data:
        jeonse, wolse = _split(data, "apt")
        print(f"✅ 전세 {len(jeonse)}건 / 월세 {len(wolse)}건")
        return jeonse, wolse
    data = _fetch(OFFI_RENT_API_URL, "offiNm", "오피스텔")
    jeonse, wolse = _split(data, "offi")
    print(f"✅ 오피스텔 전세 {len(jeonse)}건 / 월세 {len(wolse)}건")
    return jeonse, wolse


# ========================================================
# 주변 단지 실거래가
# ========================================================
def find_nearby_apts(base_lat, base_lon, apt_list, exclude_name, radius=1000, top_n=5):
    print(f"주변 단지 좌표 조회 중 (반경 {radius}m)...")
    nearby = []
    exclude_clean = exclude_name.replace(" ", "")
    for apt in apt_list:
        if exclude_clean in apt["kaptName"].replace(" ", ""):
            continue
        detail = get_apt_detail(apt["kaptCode"])
        if not detail or not detail.get("doroJuso"):
            continue
        lat2, lon2, _, jibun, umd_nm = address_to_info(detail["doroJuso"])
        if not lat2:
            continue
        dist = haversine(base_lat, base_lon, lat2, lon2)
        if dist <= radius:
            nearby.append({
                "kaptCode": apt["kaptCode"],
                "kaptName": apt["kaptName"],
                "lat":      lat2,
                "lon":      lon2,
                "dist":     dist,
                "detail":   detail,
                "jibun":    jibun or "",
                "umd_nm":   umd_nm or "",
            })
        time.sleep(0.05)

    nearby.sort(key=lambda x: x["dist"])
    print(f"✅ 반경 {radius}m 내 {len(nearby)}개 단지")
    for n in nearby[:top_n]:
        print(f"  - {n['kaptName']} ({n['dist']:.0f}m)")
    return nearby[:top_n]


def find_nearby_offi(base_lat, base_lon, radius=1000, top_n=5):
    """
    오피스텔 전용 — 카카오 키워드 검색으로 주변 오피스텔 찾기
    (국토부 단지 목록 API는 아파트 전용이라 오피스텔 지원 안 함)
    반환: [{"kaptName": ..., "lat": ..., "lon": ..., "dist": ...}, ...]
    """
    headers = {"Authorization": f"KakaoAK {KAKAO_API_KEY}"}
    # 카카오 카테고리에 오피스텔 없으므로 키워드 검색 사용
    kw_url = "https://dapi.kakao.com/v2/local/search/keyword.json"
    kw_params = {
        "query": "오피스텔",
        "x": base_lon, "y": base_lat,
        "radius": radius, "size": 15, "sort": "distance"
    }
    # 주차장, 상가, 생활지원센터 등 부속시설 제외
    EXCLUDE_SUFFIX = ["주차장", "상가", "생활지원", "관리사무소", "어린이집", "유치원",
                      "경비실", "커뮤니티", "부대시설"]

    nearby = []
    try:
        res = session.get(kw_url, headers=headers, params=kw_params, timeout=10)
        docs = res.json().get("documents", [])
        for d in docs:
            name = d.get("place_name", "")
            # 부속시설 제외
            if any(ex in name for ex in EXCLUDE_SUFFIX):
                continue
            lat2 = float(d.get("y", 0))
            lon2 = float(d.get("x", 0))
            dist = haversine(base_lat, base_lon, lat2, lon2)
            if dist <= radius:
                nearby.append({
                    "kaptName": name,
                    "lat":      lat2,
                    "lon":      lon2,
                    "dist":     dist,
                    "jibun":    "",
                    "umd_nm":   "",
                })
        nearby.sort(key=lambda x: x["dist"])
        print(f"✅ 반경 {radius}m 내 주변 오피스텔 {len(nearby)}개 발견")
    except Exception as e:
        print(f"❌ 주변 오피스텔 검색 오류: {e}")
    return nearby[:top_n]


def fetch_nearby_sales(base_lat, base_lon, apt_list, lawd_cd, exclude_name, radius=1000, months=36):
    nearby = find_nearby_apts(base_lat, base_lon, apt_list, exclude_name, radius)
    if not nearby:
        return {}

    now = datetime.now()
    print(f"\n구 전체 거래 데이터 로딩 중 ({months}개월)...")
    all_trades = {}
    lock = threading.Lock()

    def _fetch_one(ym):
        try:
            resp = gov_session.get(SALES_API_URL, params={
                "serviceKey": GOV_SERVICE_KEY,
                "LAWD_CD": lawd_cd, "DEAL_YMD": ym,
                "pageNo": "1", "numOfRows": "9999"
            }, headers=HDR_XML, timeout=30)
            if resp.status_code == 200 and resp.content:
                root = ET.fromstring(resp.content)
                for item in root.findall(".//item"):
                    row = {c.tag: (c.text.strip() if c.text else "") for c in item}
                    apt_nm = row.get("aptNm", "").replace(" ", "")
                    with lock:
                        all_trades.setdefault(apt_nm, []).append(row)
        except Exception as e:
            print(f"  [오류] {ym}: {e}")

    ym_list = [(now - relativedelta(months=i)).strftime("%Y%m") for i in range(months)]
    with ThreadPoolExecutor(max_workers=4) as pool:
        pool.map(_fetch_one, ym_list)

    # 주변 단지 jibun 목록 구성
    nearby_jibun = {
        apt["kaptName"]: {
            "jibun":  apt.get("jibun", ""),
            "umd_nm": apt.get("umd_nm", ""),
            "dist":   apt["dist"],
            "lat":    apt["lat"],
            "lon":    apt["lon"],
        }
        for apt in nearby
    }

    # RTMS 데이터를 jibun 기준으로도 인덱싱
    trades_by_jibun = {}   # key: "umdNm|jibun"
    for apt_nm, rows in all_trades.items():
        if rows:
            r0 = rows[0]
            key = f"{r0.get('umdNm','').strip()}|{r0.get('jibun','').strip()}"
            trades_by_jibun[key] = rows

    nearby_result = {}
    for apt_name, info in nearby_jibun.items():
        jibun   = info["jibun"]
        umd_nm  = info["umd_nm"]
        matched_rows = None

        # 1순위: jibun 정확 매칭
        if jibun and umd_nm:
            key = f"{umd_nm}|{jibun}"
            if key in trades_by_jibun:
                matched_rows = trades_by_jibun[key]

        # 2순위: aptNm 문자열 매칭 (fallback)
        if matched_rows is None:
            matched_key, score = _match_apt_name(apt_name, all_trades.keys())
            if matched_key and score >= 2:
                apt_clean = apt_name.replace(" ", "")
                if matched_key in apt_clean or score >= 4:
                    matched_rows = all_trades[matched_key]

        if matched_rows:
            nearby_result[apt_name] = {
                "trades": matched_rows,
                "dist":   info["dist"],
                "lat":    info["lat"],
                "lon":    info["lon"],
                "jibun":  info["jibun"],
                "umd_nm": info["umd_nm"],
            }
            print(f"  ✅ {apt_name}: {len(matched_rows)}건 ({info['dist']:.0f}m)")
        else:
            print(f"  ⚠️ {apt_name}: 거래 없음 ({info['dist']:.0f}m)")
    return nearby_result


def fetch_nearby_rent(nearby_sales, lawd_cd, months=36):
    """주변 단지 전세 데이터 조회 (nearby_sales와 동일 단지 목록 재사용)"""
    if not nearby_sales:
        return {}

    now = datetime.now()
    print(f"\n주변 단지 전세 데이터 로딩 중 ({months}개월)...")
    all_trades = {}
    lock = threading.Lock()

    def _fetch_one(ym):
        try:
            resp = gov_session.get(RENT_API_URL, params={
                "serviceKey": GOV_SERVICE_KEY,
                "LAWD_CD": lawd_cd, "DEAL_YMD": ym,
                "pageNo": "1", "numOfRows": "9999"
            }, headers=HDR_XML, timeout=30)
            if resp.status_code == 200 and resp.content:
                root = ET.fromstring(resp.content)
                for item in root.findall(".//item"):
                    row = {c.tag: (c.text.strip() if c.text else "") for c in item}
                    # 전세만 (월세 제외)
                    if row.get("monthlyRent") and row["monthlyRent"] != "0":
                        continue
                    apt_nm = row.get("aptNm", "").replace(" ", "")
                    with lock:
                        all_trades.setdefault(apt_nm, []).append(row)
        except Exception as e:
            print(f"  [오류] {ym}: {e}")

    ym_list = [(now - relativedelta(months=i)).strftime("%Y%m") for i in range(months)]
    with ThreadPoolExecutor(max_workers=4) as pool:
        pool.map(_fetch_one, ym_list)

    # 디버그
    if all_trades:
        sample = [k for k in list(all_trades.keys())[:15] if k]
        print(f"  [디버그] 전세 API 단지명 샘플: {sample}")

    # RTMS 전세 데이터를 jibun 기준으로 인덱싱
    trades_by_jibun = {}
    for apt_nm, rows in all_trades.items():
        if rows:
            r0 = rows[0]
            key = f"{r0.get('umdNm','').strip()}|{r0.get('jibun','').strip()}"
            trades_by_jibun[key] = rows

    nearby_result = {}
    for apt_name, info in nearby_sales.items():
        # nearby_sales에 jibun 정보가 있으면 우선 사용
        jibun  = info.get("jibun", "")
        umd_nm = info.get("umd_nm", "")
        matched_rows = None

        # 1순위: jibun 정확 매칭
        if jibun and umd_nm:
            key = f"{umd_nm}|{jibun}"
            if key in trades_by_jibun:
                matched_rows = trades_by_jibun[key]

        # 2순위: aptNm 문자열 매칭 (fallback)
        if matched_rows is None:
            matched_key, score = _match_apt_name(apt_name, all_trades.keys())
            if matched_key and score >= 2:
                apt_clean = apt_name.replace(" ", "")
                if matched_key in apt_clean or score >= 4:
                    matched_rows = all_trades[matched_key]

        if matched_rows:
            nearby_result[apt_name] = {
                "trades": matched_rows,
                "dist":   info["dist"],
                "lat":    info["lat"],
                "lon":    info["lon"],
            }
            print(f"  ✅ {apt_name} 전세: {len(matched_rows)}건")
        else:
            print(f"  ⚠️ {apt_name}: 전세 거래 없음")
    return nearby_result


# ========================================================
# DB 단지 조회 (Supabase apartments — 1순위)
# ========================================================
def find_apt_from_db(apt_name, gu="", dong=""):
    """
    Supabase apartments DB에서 단지 검색
    1순위: slug 정확 매칭 (공백 제거된 이름)
    2순위: kapt_name ilike 부분 매칭
    3순위: 앞 4글자 후보 추출 → 로컬 fuzzy 매칭
    반환: detail dict (find_apt 반환 형식과 동일) 또는 None
    """
    name_clean = apt_name.replace(" ", "")
    if not name_clean:
        return None

    print(f"[DB] 단지 검색: '{apt_name}'" + (f" (구: {gu})" if gu else ""))

    def _row_to_detail(row):
        """DB row → 기존 detail dict 형식으로 변환
        kaptTopFloor / kaptBcompany / kaptTotPkCnt 는 DB에 없으므로
        kaptCode 있으면 get_apt_detail() 로 보완
        """
        hh       = row.get("households") or ""
        kapt_code = row.get("kapt_code", "")

        # 국토부 상세 API 보완 (아파트 + kaptCode 있을 때만)
        extra = {}
        if kapt_code and row.get("property_type", "apt") == "apt":
            try:
                d = get_apt_detail(kapt_code)
                if d:
                    extra = {
                        "kaptTopFloor": d.get("kaptTopFloor", ""),
                        "kaptBcompany": d.get("kaptBcompany", ""),
                        "kaptTotPkCnt": d.get("kaptTotPkCnt", ""),
                        # kaptdaCnt 도 API가 더 정확하면 덮어쓰기
                        "kaptdaCnt":    d.get("kaptdaCnt", "") or (str(int(float(hh))) if hh else ""),
                        "kaptUsedate":  d.get("kaptUsedate", "") or row.get("use_date", ""),
                        "doroJuso":     d.get("doroJuso", "") or row.get("doro_juso", ""),
                    }
                    print(f"  ✅ [DB+API] 최고층:{extra['kaptTopFloor']} / 시공사:{extra['kaptBcompany']}")
            except Exception as e:
                print(f"  ⚠️ 상세 API 보완 실패: {e}")

        return {
            "kaptCode":       kapt_code,
            "kaptName":       row.get("kapt_name", ""),
            "doroJuso":       extra.get("doroJuso") or row.get("doro_juso", ""),
            "kaptdaCnt":      extra.get("kaptdaCnt") or (str(int(float(hh))) if hh else ""),
            "kaptUsedate":    extra.get("kaptUsedate") or row.get("use_date", ""),
            "kaptTopFloor":   extra.get("kaptTopFloor", "") or row.get("top_floor", ""),
            "kaptBcompany":   extra.get("kaptBcompany", "") or row.get("builder", ""),
            "kaptTotPkCnt":   extra.get("kaptTotPkCnt", "") or row.get("parking", ""),
            "hhldCount":      extra.get("kaptdaCnt") or (str(int(float(hh))) if hh else ""),
            "lat":            float(row.get("lat") or 0),
            "lon":            float(row.get("lon") or 0),
            "bjd_code":       row.get("bjd_code", ""),
            "apt_list":       [],
            "_db_prop_type":  row.get("property_type", "apt"),
            "_db_pyeongs":    row.get("pyeongs") or [],  # 공급면적 직접 사용
        }

    try:
        # ── 공통 필터 파라미터 ──
        gu_filter = {"sgg": f"eq.{gu}"} if gu else {}

        # 1순위: slug 정확 매칭
        res = requests.get(
            f"{SUPABASE_URL}/rest/v1/apartments",
            headers=SUPABASE_HEADERS,
            params={"select": "*", "slug": f"eq.{name_clean}", **gu_filter},
            timeout=10,
        )
        rows = res.json() if res.status_code == 200 else []
        if isinstance(rows, list) and rows:
            row = rows[0]
            print(f"  ✅ [DB-slug] '{row['kapt_name']}' ({row.get('doro_juso', '')})")
            detail = _row_to_detail(row)
            if detail["bjd_code"] and detail["_db_prop_type"] == "apt":
                detail["apt_list"] = get_apt_list(detail["bjd_code"])
            return detail

        # 2순위: kapt_name ilike 부분 매칭
        res2 = requests.get(
            f"{SUPABASE_URL}/rest/v1/apartments",
            headers=SUPABASE_HEADERS,
            params={"select": "*", "kapt_name": f"ilike.*{name_clean}*", **gu_filter},
            timeout=10,
        )
        rows2 = res2.json() if res2.status_code == 200 else []
        if isinstance(rows2, list) and rows2:
            # 여러 개면 fuzzy로 가장 유사한 것 선택
            best_row = rows2[0]
            best_score = 0
            for r in rows2:
                s = _fuzzy_score(name_clean, r["kapt_name"].replace(" ", ""))
                if s > best_score:
                    best_score = s
                    best_row = r
            print(f"  ✅ [DB-ilike] '{best_row['kapt_name']}' (유사도: {best_score:.2f})")
            detail = _row_to_detail(best_row)
            if detail["bjd_code"] and detail["_db_prop_type"] == "apt":
                detail["apt_list"] = get_apt_list(detail["bjd_code"])
            return detail

        # 3순위: 앞 4글자로 후보 추출 → 로컬 fuzzy 매칭
        prefix = name_clean[:4]
        res3 = requests.get(
            f"{SUPABASE_URL}/rest/v1/apartments",
            headers=SUPABASE_HEADERS,
            params={"select": "*", "kapt_name": f"ilike.*{prefix}*", **gu_filter},
            timeout=10,
        )
        candidates = res3.json() if res3.status_code == 200 else []
        if isinstance(candidates, list) and candidates:
            # _fuzzy_match_best 호환 형식으로 변환 (kaptName 키 필요)
            compat = [{"kaptName": r["kapt_name"], "_row": r} for r in candidates]
            best, score = _fuzzy_match_best(name_clean, compat, min_score=0.72)
            if best:
                row = best["_row"]
                print(f"  ✅ [DB-fuzzy] '{row['kapt_name']}' (유사도: {score:.2f})")
                detail = _row_to_detail(row)
                if detail["bjd_code"] and detail["_db_prop_type"] == "apt":
                    detail["apt_list"] = get_apt_list(detail["bjd_code"])
                return detail
            else:
                top3 = sorted(candidates,
                    key=lambda r: _fuzzy_score(name_clean, r["kapt_name"].replace(" ", "")),
                    reverse=True)[:3]
                print(f"  ⚠️ [DB] fuzzy 매칭 실패. 유사 후보:")
                for r in top3:
                    s = _fuzzy_score(name_clean, r["kapt_name"].replace(" ", ""))
                    print(f"      '{r['kapt_name']}' (유사도: {s:.2f})")

        print(f"  ℹ️  [DB] '{apt_name}' 미발견 → 기존 API 폴백")
        return None

    except Exception as e:
        print(f"  ⚠️ [DB] 조회 오류: {e} → 기존 API 폴백")
        return None


# ========================================================
# 단지 조회 메인
# ========================================================
def find_apt(parsed):
    gu_name      = parsed.get("gu", "")
    dong_name    = parsed.get("dong", "")
    apt_name     = parsed.get("apt_name", "")
    road_address = parsed.get("road_address", "")
    lat, lon, bjd_code = None, None, None

    if road_address:
        print(f"도로명 주소로 조회: {road_address}")
        lat, lon, bjd_code, _, _ = address_to_info(road_address)
        if lat:
            print(f"✅ 좌표: {lat}, {lon} / bjdCode: {bjd_code}")

    if not bjd_code:
        print(f"카카오 키워드 검색: {gu_name} {dong_name} {apt_name}")
        doc = search_kakao_keyword(f"{gu_name} {dong_name} {apt_name}")
        if doc:
            lat2, lon2, bjd_code, _, _ = address_to_info(doc.get("address_name", ""))
            if not lat:
                lat, lon = lat2, lon2

    if not bjd_code:
        print("❌ 위치를 찾을 수 없습니다.")
        return None

    apt_list = get_apt_list(bjd_code)
    if not apt_list:
        print("❌ 단지 목록 조회 실패")
        return None

    search_name = apt_name.replace(" ", "")

    # 1순위: 정확 포함 매칭
    matched = [a for a in apt_list
               if search_name in a["kaptName"].replace(" ", "")
               or a["kaptName"].replace(" ", "") in search_name]

    # 다중 매칭 시 도로명으로 2차 필터 (짧은 단지명이 여러 단지에 포함될 때)
    if len(matched) > 1 and road_address:
        road_nm_f, road_bonbun_f = _extract_road_info(road_address)
        if road_nm_f:
            refined = []
            for apt in matched:
                d = get_apt_detail(apt["kaptCode"])
                if d and d.get("doroJuso"):
                    d_road, d_bonbun = _extract_road_info(d["doroJuso"])
                    if road_nm_f in d_road or d_road in road_nm_f:
                        if not road_bonbun_f or not d_bonbun or road_bonbun_f == d_bonbun:
                            refined.append(apt)
                time.sleep(0.05)
            if len(refined) == 1:
                print(f"  ✅ 도로명으로 다중매칭 해소: '{refined[0]['kaptName']}'")
                matched = refined
            elif len(refined) > 1:
                print(f"  ⚠️ 도로명 필터 후에도 {len(refined)}개 — 가장 유사한 것 선택")
                matched = sorted(refined,
                    key=lambda a: _fuzzy_score(search_name, a["kaptName"].replace(" ","")),
                    reverse=True)[:1]

    # 2순위: 퍼지 매칭 (오타/브랜드명 차이 허용)
    if not matched:
        print(f"  정확 매칭 실패 → 퍼지 매칭 시도 (입력: '{search_name}')")
        best, score = _fuzzy_match_best(search_name, apt_list, min_score=0.82)
        if best:
            print(f"  ✅ 퍼지 매칭: '{best['kaptName']}' (유사도: {score:.2f})")
            matched = [best]
        else:
            # 상위 후보 표시 (디버그)
            top3 = sorted(apt_list, key=lambda a: _fuzzy_score(search_name, a["kaptName"]), reverse=True)[:3]
            print(f"  ⚠️ 퍼지 매칭 실패. 유사 후보:")
            for t in top3:
                s = _fuzzy_score(search_name, t["kaptName"])
                print(f"      '{t['kaptName']}' (유사도: {s:.2f})")

    # 3순위: 좌표 기준 최근접 단지 (마지막 수단 — 조건 엄격)
    if not matched and lat and lon:
        print("단지명+퍼지 매칭 실패 → 좌표 기준 최근접 단지 검색 (150m 이내)...")
        closest, min_dist = None, float("inf")
        for apt in apt_list[:20]:
            detail = get_apt_detail(apt["kaptCode"])
            if detail and detail.get("doroJuso"):
                a_lat, a_lon, _, _, _ = address_to_info(detail["doroJuso"])
                if a_lat:
                    dist = haversine(lat, lon, a_lat, a_lon)
                    if dist < min_dist:
                        min_dist = dist
                        closest = {**apt, "dist": dist}
            time.sleep(0.05)
        if closest and closest.get("dist", 9999) < 150:
            # 이름 유사도 최소 체크 — 완전 엉뚱한 단지 방지
            coord_score = _fuzzy_score(search_name, closest["kaptName"])
            if coord_score >= 0.65:
                print(f"  ✅ 좌표 매칭: '{closest['kaptName']}' ({closest['dist']:.0f}m, 유사도: {coord_score:.2f})")
                matched = [closest]
            else:
                print(f"  ⚠️ 좌표 최근접 '{closest['kaptName']}' ({closest['dist']:.0f}m) — 이름 유사도 {coord_score:.2f} 너무 낮아 탈락")
        else:
            print("❌ 150m 이내 단지 없음")
            return None

    if not matched:
        return None

    detail = get_apt_detail(matched[0]["kaptCode"])
    if detail:
        detail["lat"]      = lat
        detail["lon"]      = lon
        detail["bjd_code"] = bjd_code
        detail["apt_list"] = apt_list
        print(f"\n{'='*40}")
        print(f"단지명:   {detail['kaptName']}")
        print(f"도로명:   {detail['doroJuso']}")
        hh = detail.get('kaptdaCnt', '')
        hh_str = str(int(float(hh))) if hh else "정보없음"
        print(f"세대수:   {hh_str}세대")
        print(f"입주일:   {detail['kaptUsedate']}")
        print(f"{'='*40}")
    return detail


def find_apt_offi(parsed):
    """
    오피스텔 전용 단지 조회
    국토부 단지목록 API 대신 카카오 키워드 검색으로 기본 정보 구성
    반환: detail dict (아파트 detail과 동일 키, 없는 필드는 빈값)
    """
    gu_name      = parsed.get("gu", "")
    dong_name    = parsed.get("dong", "")
    apt_name     = parsed.get("apt_name", "")
    road_address = parsed.get("road_address", "")

    lat, lon, bjd_code = None, None, None

    # 1. 도로명 주소로 좌표/bjd_code 추출
    if road_address:
        print(f"[오피스텔] 도로명 주소로 조회: {road_address}")
        lat, lon, bjd_code, jibun, umd_nm = address_to_info(road_address)
        if lat:
            print(f"✅ 좌표: {lat}, {lon} / bjdCode: {bjd_code}")

    # 2. 실패 시 카카오 키워드 검색
    kakao_road = ""
    if not lat:
        query = f"{gu_name} {dong_name} {apt_name}".strip()
        print(f"[오피스텔] 카카오 검색: {query}")
        doc = search_kakao_keyword(query, allow_offi=True)
        if doc:
            kakao_road = doc.get("road_address_name") or doc.get("address_name", "")
            lat, lon, bjd_code, jibun, umd_nm = address_to_info(kakao_road)
            if lat:
                print(f"✅ 좌표: {lat}, {lon} / bjdCode: {bjd_code}")

    if not lat or not bjd_code:
        print("❌ [오피스텔] 위치를 찾을 수 없습니다.")
        return None

    # 3. detail 딕셔너리 구성 (오피스텔은 없는 필드 빈값)
    # 도로명: 입력된 road_address 우선, 없으면 카카오에서 얻은 주소 사용
    resolved_road = road_address or kakao_road

    # 3. detail 딕셔너리 구성 (오피스텔은 없는 필드 빈값)
    detail = {
        "kaptCode":     "",
        "kaptName":     apt_name,
        "doroJuso":     resolved_road,   # 도로명 주소 (건축물대장 API 호출용)
        "kaptdaCnt":    "",            # 세대수 정보 없음
        "kaptUsedate":  "",            # 준공일 없음 (실거래가 buildYear로 보완)
        "kaptTopFloor": "",            # 최고층 없음
        "kaptBcompany": "",            # 시공사 없음
        "kaptTotPkCnt": "",            # 주차 없음
        "hhldCount":    "",
        "lat":          lat,
        "lon":          lon,
        "bjd_code":     bjd_code,
        "apt_list":     [],            # 주변 아파트 목록 없음 (오피스텔)
    }

    print(f"\n{'='*40}")
    print(f"단지명:   {detail['kaptName']}")
    print(f"도로명:   {detail['doroJuso']}")
    print(f"유형:     오피스텔")
    print(f"{'='*40}")
    return detail


# ========================================================
# 이미지 생성 (Selenium + 카카오맵 SDK)
# ========================================================
from selenium import webdriver
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

CAPTURE_W = 1200
CAPTURE_H = 630
WINDOW_W  = 1280
WINDOW_H  = 900


def _get_driver():
    """Selenium 드라이버 생성"""
    opts = ChromeOptions()
    opts.add_argument("--log-level=3")
    opts.add_argument("--disable-blink-features=AutomationControlled")
    opts.add_experimental_option("excludeSwitches", ["enable-automation", "enable-logging"])
    opts.add_experimental_option("useAutomationExtension", False)
    # 메모리 절약 옵션
    opts.add_argument("--headless=new")           # 화면 없이 실행 (메모리 절약)
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--disable-extensions")
    opts.add_argument("--disable-software-rasterizer")
    opts.add_argument("--js-flags=--max-old-space-size=256")
    service = ChromeService(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=opts)
    driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
    driver.set_window_size(WINDOW_W, WINDOW_H)
    return driver


def _screenshot_to_image(driver, output_path):
    """스크린샷 → 1200x630 크롭 후 저장"""
    tmp = output_path + "_tmp.png"
    driver.save_screenshot(tmp)
    img = Image.open(tmp).convert("RGB")
    W, H = img.size
    x0 = max(0, (W - CAPTURE_W) // 2)
    y0 = max(0, (H - CAPTURE_H) // 2)
    cropped = img.crop((x0, y0, x0 + CAPTURE_W, y0 + CAPTURE_H))
    cropped.save(output_path)
    img.close()
    cropped.close()
    try:
        os.remove(tmp)
    except OSError:
        pass
    return output_path


def _build_simple_map_html(lat, lon, apt_name):
    """단순지도 HTML (카카오맵 SDK)"""
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>html,body,#map{{width:100%;height:100%;margin:0;padding:0;}}</style>
</head><body><div id="map"></div>
<script src="//dapi.kakao.com/v2/maps/sdk.js?appkey={KAKAO_JS_KEY}&autoload=false"></script>
<script>
kakao.maps.load(function(){{
  var map = new kakao.maps.Map(document.getElementById('map'),{{
    center: new kakao.maps.LatLng({lat},{lon}), level:4
  }});
  var marker = new kakao.maps.Marker({{
    position: new kakao.maps.LatLng({lat},{lon}), map:map
  }});
  var overlay = new kakao.maps.CustomOverlay({{
    position: new kakao.maps.LatLng({lat},{lon}),
    content:'<div style="background:#1a1a2e;color:#f5c518;font-size:13px;font-weight:900;padding:5px 12px;border-radius:8px;box-shadow:0 2px 6px rgba(0,0,0,0.4);">{apt_name}</div>',
    yAnchor:2.8, map:map
  }});
}});
</script></body></html>"""


def _build_label_map_html(base_lat, base_lon, apt_name, markers, analysis_type):
    """라벨지도 HTML — 최대 거리 기반 동적 줌 (참조: 통합_v8_3)"""
    color = "#4A90E2" if analysis_type == "전세가" else "#F78148"

    # ── 줌 레벨 — 가장 먼 마커 거리 기준 ──────────────
    if not markers:
        level = 4   # 단독: 동네 수준 (level 5는 너무 넓음)
    else:
        max_dist = max(m["dist"] for m in markers) / 1000  # km 단위
        if max_dist <= 0.5:
            level = 3   # 500m 이내: 최대 확대
        elif max_dist <= 0.8:
            level = 4   # 800m 이내
        elif max_dist <= 1.2:
            level = 5   # 1.2km 이내
        else:
            level = 6   # 1.2km 초과
    # ────────────────────────────────────────────────────


    OFFI_COLOR = "#4A90E2"
    APT_COLOR  = "#F78148"
    TAIL = 40
    DIRECTIONS = ["up","down","right","left","up-right","up-left","down-right","down-left"]

    # 기준 단지 오버레이 — 흰 배경 + 빨간 테두리 (주변 마커와 같은 구조, 색으로 차별화)
    base_js = (
        "  (function(){"
        "var wrap=document.createElement('div');"
        "wrap.style.cssText='display:flex;flex-direction:column;align-items:center;';"
        "var box=document.createElement('div');"
        "box.style.cssText='background:white;padding:5px 14px;"
        "border-radius:10px;white-space:nowrap;"
        "box-shadow:0 2px 8px rgba(0,0,0,0.3);border:3px solid #E8303A;"
        "text-align:center;line-height:1.5;';"
        "var nm=document.createElement('div');"
        f"nm.innerText='{apt_name}';"
        "nm.style.cssText='font-size:13px;color:#E8303A;font-weight:900;';"
        "box.appendChild(nm);"
        "var line=document.createElement('div');"
        "line.style.cssText='width:2px;height:60px;background:#E8303A;';"
        "var arr=document.createElement('div');"
        "arr.style.cssText='width:0;height:0;border-left:6px solid transparent;"
        "border-right:6px solid transparent;border-top:8px solid #E8303A;';"
        "wrap.appendChild(box);wrap.appendChild(line);wrap.appendChild(arr);"
        f"new kakao.maps.CustomOverlay({{position:new kakao.maps.LatLng({base_lat},{base_lon}),"
        "content:wrap,yAnchor:1.0,zIndex:10,map:map});"
        "})();"
    )

    # ── 지도 경계 추정 (줌 레벨 기반 lat/lon 범위) ──────
    # 카카오맵 level별 1px당 도 단위 (대략값, 서울 위도 기준)
    ZOOM_DEG = {3: 0.00003, 4: 0.00006, 5: 0.00012, 6: 0.00024}
    deg_per_px = ZOOM_DEG.get(level, 0.0001)
    MAP_W, MAP_H = 1200, 630   # 캡처 해상도
    # 지도 중심 = base 단지
    map_lat_min = float(base_lat) - deg_per_px * (MAP_H / 2)
    map_lat_max = float(base_lat) + deg_per_px * (MAP_H / 2)
    map_lon_min = float(base_lon) - deg_per_px * (MAP_W / 2) * 0.7  # 경도 보정
    map_lon_max = float(base_lon) + deg_per_px * (MAP_W / 2) * 0.7

    # 경계 여유 (전체 범위의 20%)
    lat_margin = (map_lat_max - map_lat_min) * 0.20
    lon_margin = (map_lon_max - map_lon_min) * 0.20

    def _safe_directions(mlat, mlon):
        """지도 경계를 벗어나지 않는 안전한 방향 목록 반환"""
        safe = []
        near_bottom = mlat < map_lat_min + lat_margin
        near_top    = mlat > map_lat_max - lat_margin
        near_left   = mlon < map_lon_min + lon_margin
        near_right  = mlon > map_lon_max - lon_margin

        for d in ["up", "down", "right", "left", "up-right", "up-left", "down-right", "down-left"]:
            # 경계 근처면 해당 방향 제외
            if near_bottom and "down" in d:
                continue
            if near_top and "up" in d:
                continue
            if near_left and "left" in d:
                continue
            if near_right and "right" in d:
                continue
            safe.append(d)
        return safe if safe else ["up"]  # 최후 폴백

    # 이미 배치된 말풍선 영역 추적 (겹침 방지)
    placed_boxes = []  # [(center_lat, center_lon, w_deg, h_deg), ...]
    BOX_W = deg_per_px * 140 * 0.7  # 말풍선 가로 약 140px
    BOX_H = deg_per_px * 55          # 말풍선 세로 약 55px

    def _overlaps(mlat, mlon, direction):
        """해당 방향으로 배치 시 기존 말풍선과 겹치는지 검사"""
        TAIL_DEG = deg_per_px * TAIL
        if "up" in direction:
            bx, by = mlon, mlat + TAIL_DEG + BOX_H / 2
        elif "down" in direction:
            bx, by = mlon, mlat - TAIL_DEG - BOX_H / 2
        elif direction == "right":
            bx, by = mlon + TAIL_DEG + BOX_W / 2, mlat
        else:  # left
            bx, by = mlon - TAIL_DEG - BOX_W / 2, mlat

        for (px, py, pw, ph) in placed_boxes:
            if (abs(bx - px) < (BOX_W + pw) / 2 and
                    abs(by - py) < (BOX_H + ph) / 2):
                return True
        return False

    def _pick_direction(mlat, mlon, dlat, dlon, dist_m):
        """안전하고 겹치지 않는 최적 방향 선택"""
        safe = _safe_directions(mlat, mlon)

        # 1순위: 기준 단지에서 멀어지는 방향 (자연스럽게 연결)
        preferred = []
        if dlat >= 0 and dlon >= 0:
            preferred = ["up-right", "up", "right"]
        elif dlat >= 0 and dlon < 0:
            preferred = ["up-left", "up", "left"]
        elif dlat < 0 and dlon >= 0:
            preferred = ["down-right", "down", "right"]
        else:
            preferred = ["down-left", "down", "left"]

        # 먼 단지(500m+)는 4방향 우선
        if dist_m >= 500:
            abs_lat, abs_lon = abs(dlat), abs(dlon)
            if abs_lat >= abs_lon:
                preferred = (["up"] if dlat >= 0 else ["down"]) + preferred
            else:
                preferred = (["right"] if dlon >= 0 else ["left"]) + preferred

        # safe 안에서 겹치지 않는 방향 선택
        for d in preferred:
            if d in safe and not _overlaps(mlat, mlon, d):
                return d
        # 겹쳐도 safe 방향 중 첫번째
        for d in preferred:
            if d in safe:
                return d
        return safe[0]

    # 주변 마커 오버레이 — 실제 좌표 기반 방향 계산
    markers_js = ""
    for i, m in enumerate(markers):
        price_str = format_price_word(m["price"])
        mtype     = m.get("marker_type", "")
        mc        = OFFI_COLOR if mtype == "offi" else (APT_COLOR if mtype == "apt" else color)
        mlat, mlon = float(m["lat"]), float(m["lon"])

        dlat = mlat - float(base_lat)
        dlon = mlon - float(base_lon)
        dist_m = m.get("dist", 500)

        # 거리 기반 꼭짓점 길이
        if dist_m < 100:
            TAIL = 100
        elif dist_m < 300:
            TAIL = 69
        elif dist_m < 500:
            TAIL = 44
        else:
            TAIL = 25

        # 방향 결정 (경계 + 겹침 고려)
        direction = _pick_direction(mlat, mlon, dlat, dlon, dist_m)

        # 배치된 말풍선 등록
        TAIL_DEG = deg_per_px * TAIL
        if "up" in direction:
            bx, by = mlon, mlat + TAIL_DEG + BOX_H / 2
        elif "down" in direction:
            bx, by = mlon, mlat - TAIL_DEG - BOX_H / 2
        elif direction == "right":
            bx, by = mlon + TAIL_DEG + BOX_W / 2, mlat
        else:
            bx, by = mlon - TAIL_DEG - BOX_W / 2, mlat
        placed_boxes.append((bx, by, BOX_W, BOX_H))

        # 말풍선 텍스트: 단지명 + 공급면적 + 가격
        marker_name = m.get("name", "")
        area_str    = m.get("area_str", "")

        # 말풍선 방향별 CSS 설정
        if direction == "up":
            flex = "column"
            inner = "wrap.appendChild(box);wrap.appendChild(line);wrap.appendChild(arr);"
            line_css = f"width:2px;height:{TAIL}px;background:{mc};"
            arr_css  = f"width:0;height:0;border-left:6px solid transparent;border-right:6px solid transparent;border-top:8px solid {mc};"
            y_anc = 1.0
        elif direction == "down":
            flex = "column"
            inner = "wrap.appendChild(arr);wrap.appendChild(line);wrap.appendChild(box);"
            line_css = f"width:2px;height:{TAIL}px;background:{mc};"
            arr_css  = f"width:0;height:0;border-left:6px solid transparent;border-right:6px solid transparent;border-bottom:8px solid {mc};"
            y_anc = 0.0
        elif direction == "right":
            flex = "row"
            inner = "wrap.appendChild(arr);wrap.appendChild(line);wrap.appendChild(box);"
            line_css = f"height:2px;width:{TAIL}px;background:{mc};"
            arr_css  = f"width:0;height:0;border-top:6px solid transparent;border-bottom:6px solid transparent;border-right:8px solid {mc};"
            y_anc = 0.5
        elif direction == "left":
            flex = "row"
            inner = "wrap.appendChild(box);wrap.appendChild(line);wrap.appendChild(arr);"
            line_css = f"height:2px;width:{TAIL}px;background:{mc};"
            arr_css  = f"width:0;height:0;border-top:6px solid transparent;border-bottom:6px solid transparent;border-left:8px solid {mc};"
            y_anc = 0.5
        elif direction == "up-right":
            # 대각선: 말풍선 우상단, 꼭짓점 좌하 방향
            flex = "column"
            inner = "wrap.appendChild(box);wrap.appendChild(line);wrap.appendChild(arr);"
            line_css = f"width:2px;height:{TAIL}px;background:{mc};margin-left:{TAIL//2}px;"
            arr_css  = f"width:0;height:0;border-left:6px solid transparent;border-right:6px solid transparent;border-top:8px solid {mc};margin-left:{TAIL//2}px;"
            y_anc = 1.0
        elif direction == "up-left":
            # 대각선: 말풍선 좌상단, 꼭짓점 우하 방향
            flex = "column"
            inner = "wrap.appendChild(box);wrap.appendChild(line);wrap.appendChild(arr);"
            line_css = f"width:2px;height:{TAIL}px;background:{mc};margin-right:{TAIL//2}px;"
            arr_css  = f"width:0;height:0;border-left:6px solid transparent;border-right:6px solid transparent;border-top:8px solid {mc};margin-right:{TAIL//2}px;"
            y_anc = 1.0
        elif direction == "down-right":
            # 대각선: 말풍선 우하단, 꼭짓점 좌상 방향
            flex = "column"
            inner = "wrap.appendChild(arr);wrap.appendChild(line);wrap.appendChild(box);"
            line_css = f"width:2px;height:{TAIL}px;background:{mc};margin-left:{TAIL//2}px;"
            arr_css  = f"width:0;height:0;border-left:6px solid transparent;border-right:6px solid transparent;border-bottom:8px solid {mc};margin-left:{TAIL//2}px;"
            y_anc = 0.0
        else:  # down-left
            # 대각선: 말풍선 좌하단, 꼭짓점 우상 방향
            flex = "column"
            inner = "wrap.appendChild(arr);wrap.appendChild(line);wrap.appendChild(box);"
            line_css = f"width:2px;height:{TAIL}px;background:{mc};margin-right:{TAIL//2}px;"
            arr_css  = f"width:0;height:0;border-left:6px solid transparent;border-right:6px solid transparent;border-bottom:8px solid {mc};margin-right:{TAIL//2}px;"
            y_anc = 0.0

        markers_js += (
            f"  (function(){{"
            f"var wrap=document.createElement('div');"
            f"wrap.style.cssText='display:flex;flex-direction:{flex};align-items:center;';"
            f"var box=document.createElement('div');"
            f"box.style.cssText='background:white;color:#222;padding:5px 11px;"
            f"border-radius:10px;font-size:12px;font-weight:bold;white-space:nowrap;"
            f"box-shadow:0 2px 5px rgba(0,0,0,0.25);border:2px solid {mc};"
            f"text-align:center;line-height:1.5;';"
            f"var nm=document.createElement('div');"
            f"nm.innerText='{marker_name}';"
            f"nm.style.cssText='font-size:11px;color:#666;font-weight:normal;';"
            f"var ar=document.createElement('div');"
            f"ar.innerText='{area_str}';"
            f"ar.style.cssText='font-size:11px;color:#888;font-weight:normal;';"
            f"var pr=document.createElement('div');"
            f"pr.innerText='{price_str}';"
            f"pr.style.cssText='font-size:13px;color:#222;font-weight:bold;';"
            f"box.appendChild(nm);{'box.appendChild(ar);' if area_str else ''}box.appendChild(pr);"
            f"var line=document.createElement('div');"
            f"line.style.cssText='{line_css}';"
            f"var arr=document.createElement('div');"
            f"arr.style.cssText='{arr_css}';"
            f"{inner}"
            f"new kakao.maps.CustomOverlay({{position:new kakao.maps.LatLng({mlat},{mlon}),"
            f"content:wrap,yAnchor:{y_anc},zIndex:1,map:map}});"
            f"}})();\n"
        )

    overlays_js = base_js + "\n" + markers_js

    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>html,body,#map{{width:100%;height:100%;margin:0;padding:0;}}</style>
</head><body><div id="map"></div>
<script src="//dapi.kakao.com/v2/maps/sdk.js?appkey={KAKAO_JS_KEY}&autoload=false"></script>
<script>
kakao.maps.load(function(){{
  var map = new kakao.maps.Map(document.getElementById('map'),{{
    center: new kakao.maps.LatLng({base_lat},{base_lon}), level:{level}
  }});
  {overlays_js}
}});
</script></body></html>"""


def _capture_html(driver, html_content, output_path, port, wait_sec=4):
    """HTML → 임시 파일 → localhost HTTP 서버로 접근 → 캡처"""
    html_filename = os.path.basename(output_path) + "_map.html"
    html_dir = os.path.dirname(output_path)
    html_path = os.path.join(html_dir, html_filename)
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    driver.get(f"http://localhost:{port}/{html_filename}")
    try:
        WebDriverWait(driver, 15).until(
            EC.presence_of_element_located((By.ID, "map"))
        )
    except Exception as e:
        print(f"  ⚠️ 지도 로딩 대기 타임아웃: {e}")
    time.sleep(wait_sec)
    _screenshot_to_image(driver, output_path)
    try:
        os.remove(html_path)
    except OSError:
        pass


def get_map_image(driver, lat, lon, apt_name, output_dir, port, seo_name=None):
    """1. 단순지도"""
    print("  단순 지도 생성 중...")
    fname = f"{seo_name}.png" if seo_name else f"{apt_name}_단순지도.png"
    path  = os.path.join(output_dir, fname)
    html  = _build_simple_map_html(lat, lon, apt_name)
    _capture_html(driver, html, path, port)
    print(f"  ✅ 단순지도 저장: {path}")
    return path


def get_label_map(driver, lat, lon, apt_name, nearby_sales, price_type, output_dir, port, seo_name=None):
    """2/3. 매매/전세 라벨지도"""
    label = "매매가" if price_type == "deal" else "전세가"
    print(f"  {label} 라벨지도 생성 중...")

    markers = []
    for name, info in nearby_sales.items():
        trades = info.get("trades", [])
        if not trades:
            continue

        # 날짜 기준 정렬
        sorted_trades = sorted(
            trades,
            key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2), str(x.get("dealDay","")).zfill(2)),
            reverse=True
        )

        # 가장 많이 거래된 평수 선택 → 동률이면 전용면적 중간값 평수 선택
        pyeong_counts = Counter(
            get_pyeong(t.get("excluUseAr"), t.get("buildYear"))
            for t in trades
            if get_pyeong(t.get("excluUseAr"), t.get("buildYear")) is not None
        )
        if pyeong_counts:
            max_cnt = pyeong_counts.most_common(1)[0][1]
            top_pyeongs = [py for py, cnt in pyeong_counts.items() if cnt == max_cnt]
            if len(top_pyeongs) == 1:
                main_pyeong = top_pyeongs[0]
            else:
                # 동률 → 전용면적 기준 중간값 평수 선택
                top_pyeongs_sorted = sorted(top_pyeongs)
                main_pyeong = top_pyeongs_sorted[len(top_pyeongs_sorted) // 2]
        else:
            main_pyeong = None
        target_trades = (
            [t for t in sorted_trades
             if get_pyeong(t.get("excluUseAr"), t.get("buildYear")) == main_pyeong]
            if main_pyeong else sorted_trades
        )

        # ── [DEBUG] 평수별 거래 현황 + 선택 근거 출력 ──────
        _debug = os.environ.get("HWIK_DEBUG")
        if _debug:
            print(f"\n  ╔═ DEBUG [{label}] {name} ════════════════════════")
            print(f"  ║ 총 거래건수: {len(trades)}건")
            print(f"  ║ 평수별 거래수:")
            for py, cnt in pyeong_counts.most_common():
                marker = " ◀ 선택" if py == main_pyeong else ""
                # 해당 평수의 전용면적 샘플 추출
                sample_ar = next(
                    (t.get("excluUseAr","?") for t in trades
                     if get_pyeong(t.get("excluUseAr"), t.get("buildYear")) == py), "?")
                print(f"  ║   {py}평 (전용 {sample_ar}㎡): {cnt}건{marker}")
            print(f"  ║ ── 주력평수({main_pyeong}평) 최근 거래 5건 ──")
            for t in target_trades[:5]:
                yr  = t.get("dealYear","")
                mo  = str(t.get("dealMonth","")).zfill(2)
                dy  = str(t.get("dealDay","")).zfill(2)
                ar  = t.get("excluUseAr","?")
                flr = t.get("floor","?")
                if price_type == "deal":
                    amt = t.get("dealAmount","?")
                    print(f"  ║   {yr}.{mo}.{dy}  전용{ar}㎡  {flr}층  → {amt}만원")
                else:
                    pre = t.get("preDeposit","")
                    dep = t.get("deposit","")
                    mrt = t.get("monthlyRent","")
                    price_disp = f"보증{pre or dep}만 / 월{mrt}만" if mrt and mrt != "0" else f"전세 {pre or dep}만원"
                    print(f"  ║   {yr}.{mo}.{dy}  전용{ar}㎡  {flr}층  → {price_disp}")
            print(f"  ╚══════════════════════════════════════════════")

        latest_price = None
        latest_trade = None  # 가격과 면적을 같은 거래에서 가져오기 위해
        for t in target_trades:
            if price_type == "deal":
                val = parse_price(t.get("dealAmount"))
            else:
                pre = t.get("preDeposit", "")
                dep = t.get("deposit", "")
                val = parse_price(pre) if _is_valid_num(pre) else parse_price(dep)
            if not val or val < 3000:
                continue
            # 매매가 단가 검증: ㎡당 400만원 미만이면 전세/임대 데이터 의심 → 제외
            if price_type == "deal":
                try:
                    exclu_f = float(str(t.get("excluUseAr","0")).replace(",","").strip())
                    if exclu_f > 0 and val / exclu_f < 400:
                        print(f"  ⚠️ [{name}] 가격 이상값 제외: {val:,}만원 / {exclu_f}㎡ = {val/exclu_f:.0f}만원/㎡")
                        continue
                except:
                    pass
            latest_price = val
            latest_trade = t  # 이 거래건의 면적을 사용
            break

        if not latest_price:
            if _debug:
                print(f"  ⚠️ [{label}] {name}: 유효 가격 없음 → 마커 제외")
            continue

        # 가격과 동일한 거래건의 전용면적 → 공급면적 변환
        main_exclu = latest_trade.get("excluUseAr", "") if latest_trade else ""
        supply_map_info = info.get("supply_map", {})
        supply_val = get_supply_from_map(main_exclu, supply_map_info) if main_exclu else None
        if supply_val:
            area_str = f"{supply_val:.0f}㎡"
        elif main_exclu:
            try:
                area_str = f"전용{float(main_exclu):.0f}㎡"
            except:
                area_str = ""
        else:
            area_str = ""

        markers.append({
            "name":     name,
            "lat":      info["lat"],
            "lon":      info["lon"],
            "dist":     info["dist"],
            "price":    latest_price,
            "area_str": area_str,
        })

    kind  = "매매가_라벨지도" if price_type == "deal" else "전세가_라벨지도"
    fname = f"{seo_name}.jpg" if seo_name else f"{apt_name}_{kind}.jpg"
    path  = os.path.join(output_dir, fname)
    html = _build_label_map_html(lat, lon, apt_name, markers, label)
    _capture_html(driver, html, path, port)
    # JPG 변환
    img = Image.open(path).convert("RGB")
    img.save(path.replace(".png", ".jpg") if path.endswith(".png") else path, "JPEG", quality=92)
    img.close()
    print(f"  ✅ {label} 라벨지도 저장: {path}")
    return path


def get_sales_chart(sales, apt_name, output_dir, jeonse=None, supply_map=None, seo_name=None):
    """4. 실거래가 추이 그래프 — 매매 + 전세 함께 표시"""
    print("  실거래가 그래프 생성 중...")
    if not sales:
        print("  ⚠️ 매매 데이터 없어 그래프 생략")
        return None

    # ── 폰트 설정 ────────────────────────────────────
    font_candidates = [
        r"C:\Windows\Fonts\malgun.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
    ]
    font_path = None
    for fp in font_candidates:
        if os.path.exists(fp):
            font_path = fp
            fm.fontManager.addfont(fp)
            plt.rcParams["font.family"] = fm.FontProperties(fname=fp).get_name()
            break
    if not font_path:
        print("  ⚠️ 한글 폰트를 찾을 수 없음")
    plt.rcParams["axes.unicode_minus"] = False

    # ── 주력 평형 선택 (거래량 기준) ────────────────────
    def _get_main_pyeong(data):
        pyeong_counts = Counter(
            get_pyeong(r.get("excluUseAr"), r.get("buildYear"))
            for r in data
            if get_pyeong(r.get("excluUseAr"), r.get("buildYear")) is not None
        )
        if not pyeong_counts:
            return None
        max_cnt = pyeong_counts.most_common(1)[0][1]
        top = [p for p, c in pyeong_counts.items() if c == max_cnt]
        return sorted(top)[len(top) // 2]

    def _build_monthly_avg(data, price_field="dealAmount", is_rent=False):
        """월별 평균가 딕셔너리 반환"""
        monthly = {}
        for r in data:
            ym = f"{r.get('dealYear','')}.{str(r.get('dealMonth','')).zfill(2)}"
            if is_rent:
                dep = r.get("deposit","")
                pre = r.get("preDeposit","")
                val = parse_price(dep if _is_valid_num(dep) else pre)
            else:
                val = parse_price(r.get(price_field))
            if val and val > 0:
                monthly.setdefault(ym, []).append(val)
        return {ym: sum(v)/len(v) for ym, v in monthly.items()}

    main_pyeong = _get_main_pyeong(sales)
    sale_data   = [r for r in sales
                   if get_pyeong(r.get("excluUseAr"), r.get("buildYear")) == main_pyeong] \
                  if main_pyeong else sales

    sale_monthly = _build_monthly_avg(sale_data)
    if not sale_monthly:
        return None

    # 전세 주력평형 (매매 주력평형과 같은 평형 우선)
    jeon_monthly = {}
    if jeonse:
        jeon_data = [r for r in jeonse
                     if get_pyeong(r.get("excluUseAr"), r.get("buildYear")) == main_pyeong]
        if not jeon_data:
            jeon_data = jeonse
        jeon_monthly = _build_monthly_avg(jeon_data, is_rent=True)

    # ── 공급면적 라벨 ────────────────────────────────
    if supply_map and main_pyeong:
        sample_exclu = next(
            (r.get("excluUseAr") for r in sale_data if r.get("excluUseAr")), None)
        supply_val = get_supply_from_map(sample_exclu, supply_map) if sample_exclu else None
        area_label = f"공급 {supply_val:.0f}㎡" if supply_val else \
                     (f"전용 {float(sample_exclu):.0f}㎡" if sample_exclu else f"{main_pyeong}평형")
    else:
        area_label = f"{main_pyeong}평형" if main_pyeong else ""

    # ── 전체 날짜 축 (매매+전세 합집합, 정렬) ─────────
    all_dates = sorted(set(list(sale_monthly.keys()) + list(jeon_monthly.keys())))
    if not all_dates:
        return None

    sale_vals = [sale_monthly.get(d) for d in all_dates]
    jeon_vals = [jeon_monthly.get(d) for d in all_dates] if jeon_monthly else []

    # x축 인덱스
    x = list(range(len(all_dates)))

    # ── 그래프 디자인 ────────────────────────────────
    fig, ax = plt.subplots(figsize=(13, 6.5))
    fig.patch.set_facecolor("#0f1923")
    ax.set_facecolor("#0f1923")

    # 그리드
    ax.grid(axis="y", color="#ffffff18", linewidth=0.8, linestyle="--")
    ax.grid(axis="x", color="#ffffff0a", linewidth=0.5, linestyle=":")

    # 색상
    SALE_COLOR = "#F5A623"   # 골드
    JEON_COLOR = "#4A90E2"   # 블루

    def _plot_line(ax, x_vals, y_vals, color, label):
        """None 제외하고 연결선 그리기"""
        valid = [(xi, yi) for xi, yi in zip(x_vals, y_vals) if yi is not None]
        if len(valid) < 2:
            return
        xs, ys = zip(*valid)
        # 영역 채우기
        y_min = min(ys) * 0.97
        ax.fill_between(xs, y_min, ys, alpha=0.12, color=color)
        # 선
        ax.plot(xs, ys, color=color, linewidth=2.5, zorder=3)
        # 점
        ax.scatter(xs, ys, color=color, s=45, zorder=4, edgecolors="#0f1923", linewidth=1.2)
        # 최근 3개 포인트에 가격 라벨
        for xi, yi in list(valid)[-3:]:
            eok  = int(yi) // 10000
            man  = int(yi) % 10000
            if eok > 0 and man > 0:
                lbl = f"{eok}억{man//100*100//100:,}"
            elif eok > 0:
                lbl = f"{eok}억"
            else:
                lbl = f"{man:,}만"
            ax.annotate(lbl,
                xy=(xi, yi), xytext=(0, 10),
                textcoords="offset points",
                ha="center", fontsize=8.5,
                color=color, fontweight="bold",
                bbox=dict(boxstyle="round,pad=0.2", fc="#0f1923", ec=color, alpha=0.7, lw=0.8))

        # 범례 끝점 라벨
        ax.annotate(f" {label}",
            xy=(xs[-1], ys[-1]),
            xytext=(6, 0), textcoords="offset points",
            va="center", fontsize=10, color=color, fontweight="bold")

    _plot_line(ax, x, sale_vals, SALE_COLOR, f"매매 ({area_label})")
    if jeon_vals:
        _plot_line(ax, x, jeon_vals, JEON_COLOR, f"전세 ({area_label})")

    # ── 축 설정 ──────────────────────────────────────
    # x축: 6개월 간격으로 레이블
    step = max(1, len(all_dates) // 8)
    ax.set_xticks(x[::step])
    ax.set_xticklabels(all_dates[::step], rotation=40, ha="right",
                       fontsize=9, color="#aaaaaa")
    ax.set_xlim(-0.5, len(all_dates) - 0.3)

    # y축: 억 단위
    def _fmt_y(val, _):
        eok = val / 10000
        if eok >= 1:
            return f"{eok:.1f}억"
        return f"{int(val):,}만"
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(_fmt_y))
    ax.tick_params(axis="y", colors="#aaaaaa", labelsize=9)
    ax.tick_params(axis="x", colors="#aaaaaa")

    # y축 여유
    all_valid = [v for v in sale_vals + jeon_vals if v]
    if all_valid:
        ax.set_ylim(min(all_valid) * 0.93, max(all_valid) * 1.07)

    # 축선 색상
    for spine in ax.spines.values():
        spine.set_edgecolor("#333333")

    # ── 제목 & 부제 ──────────────────────────────────
    fig.text(0.04, 0.95, f"{apt_name}", fontsize=15, fontweight="bold",
             color="#ffffff", va="top")
    fig.text(0.04, 0.88, "최근 3년 실거래가 추이  |  국토교통부",
             fontsize=9, color="#888888", va="top")

    # ── 워터마크 ─────────────────────────────────────
    fig.text(0.97, 0.03, "hwik.kr", fontsize=9, color="#333333",
             ha="right", va="bottom", style="italic")

    plt.tight_layout(rect=[0, 0, 0.94, 0.85])

    fname = f"{seo_name}.png" if seo_name else f"{apt_name}_실거래가그래프.png"
    path  = os.path.join(output_dir, fname)
    plt.savefig(path, dpi=120, bbox_inches="tight",
                facecolor=fig.get_facecolor(), pad_inches=0.15)
    plt.close()
    print(f"  ✅ 그래프 저장: {path}")
    return path


def build_offi_nearby_sales(base_lat, base_lon, base_sales, lawd_cd, max_total=5, base_name=""):
    """
    오피스텔 주변 시세 구성
    1순위: 반경 1km 내 오피스텔 (실거래가 있는 것)
    2순위: 오피스텔 수 < max_total 이면 → 아파트로 채움
           아파트 필터: 기준 오피스텔 최근 매매가 ± 40% 이내 우선, 없으면 거리순
    반환: nearby_sales dict (marker_type 필드 포함)
    """
    result = {}
    base_clean = base_name.replace(" ", "")  # 기준단지 제외용

    # 기준 오피스텔 최근 매매가 계산
    base_price = None
    if base_sales:
        sorted_base = sorted(base_sales,
            key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2)),
            reverse=True)
        for t in sorted_base:
            v = parse_price(t.get("dealAmount"))
            if v:
                base_price = v
                break

    # ── 1단계: 주변 오피스텔 조회 ──────────────────────
    nearby_offi = find_nearby_offi(base_lat, base_lon, radius=1000, top_n=max_total + 3)
    for offi in nearby_offi:
        nm = offi["kaptName"]
        nm_clean = nm.replace(" ", "")

        # 기준 단지 자신이면 제외 (카카오가 자기 자신을 반환하는 경우 방어)
        if base_clean and (base_clean in nm_clean or nm_clean in base_clean):
            print(f"  ⏭️ 기준단지 자신 제외: {nm}")
            continue
        trades, _ = fetch_sales_auto(nm, lawd_cd, months=12, force_offi=True)
        if trades:
            # [DEBUG] 최근 거래 상세 출력
            if os.environ.get("HWIK_DEBUG"):
                sorted_t = sorted(trades,
                    key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2)),
                    reverse=True)
                top = sorted_t[0]
                print(f"\n  ╔═ DEBUG 주변오피스텔 가격 ══════════════════")
                print(f"  ║ 단지명: {nm!r}")
                print(f"  ║ 총거래: {len(trades)}건")
                print(f"  ║ 최근거래: {top.get('dealYear')}.{top.get('dealMonth')} "
                      f"{top.get('aptNm', top.get('offiNm',''))} "
                      f"→ {top.get('dealAmount')}만원")
                print(f"  ╚══════════════════════════════════════════")
            result[nm] = {
                "trades":      trades,
                "dist":        offi["dist"],
                "lat":         offi["lat"],
                "lon":         offi["lon"],
                "jibun":       "",
                "umd_nm":      "",
                "marker_type": "offi",
            }
            print(f"  ✅ [오피스텔] {nm}: {len(trades)}건 ({offi['dist']:.0f}m)")
        if len(result) >= max_total:
            break

    # ── 2단계: 부족분 아파트로 채움 ────────────────────
    needed = max_total - len(result)
    if needed > 0:
        print(f"  오피스텔 {len(result)}개 → 아파트 {needed}개로 보완...")
        headers = {"Authorization": f"KakaoAK {KAKAO_API_KEY}"}
        # 국토부 단지 목록을 직접 활용 — bjd_code 기준으로 아파트 목록 조회
        # (AT4 카테고리는 엉뚱한 결과 반환하므로 사용 안 함)
        apt_docs = []
        try:
            # 국토부 단지 목록 (아파트만 반환하는 신뢰할 수 있는 소스)
            apt_list_raw = get_apt_list(lawd_cd + "00000"[:10-len(lawd_cd)])
            if not apt_list_raw:
                apt_list_raw = get_apt_list(lawd_cd)
            for apt in apt_list_raw:
                # 단지 상세에서 좌표 추출
                d = get_apt_detail(apt["kaptCode"])
                if not d or not d.get("doroJuso"):
                    continue
                lat2, lon2, _, _, _ = address_to_info(d["doroJuso"])
                if not lat2:
                    continue
                dist = haversine(base_lat, base_lon, lat2, lon2)
                if dist <= 1000:
                    apt_docs.append({
                        "place_name": apt["kaptName"],
                        "y": lat2, "x": lon2, "dist": dist
                    })
                time.sleep(0.05)
            apt_docs.sort(key=lambda x: x["dist"])
            apt_docs = apt_docs[:15]
            print(f"  국토부 단지 목록 기반 주변 아파트 {len(apt_docs)}개 후보")
        except Exception as e:
            print(f"  아파트 목록 조회 오류: {e}")

        # 단지명 유효성 필터 — 너무 일반적인 이름 제외
        GENERIC = {"아파트", "오피스텔", "빌라", "주상복합"}
        apt_docs = [d for d in apt_docs
                    if len(d.get("place_name","").replace(" ","")) >= 3
                    and d.get("place_name","").replace(" ","") not in GENERIC]

        apt_candidates = []
        for d in apt_docs:
            nm   = d.get("place_name", "")
            lat2 = float(d.get("y", 0))
            lon2 = float(d.get("x", 0))
            dist = haversine(base_lat, base_lon, lat2, lon2)
            if nm in result:
                continue
            trades, _ = fetch_sales_auto(nm, lawd_cd, months=12, force_offi=False)
            if not trades:
                continue
            # 최근 매매가 추출
            latest_price = None
            for t in sorted(trades,
                key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2)),
                reverse=True):
                v = parse_price(t.get("dealAmount"))
                if not v or v < 3000:
                    continue
                try:
                    exclu_f = float(str(t.get("excluUseAr","0")).replace(",","").strip())
                    if exclu_f > 0 and v / exclu_f < 400:
                        continue  # ㎡당 400만원 미만 → 전세/임대 데이터 의심
                except:
                    pass
                latest_price = v
                break
            if not latest_price:
                continue

            # 가격 유사도 점수 (기준가 대비 차이율, 작을수록 좋음)
            if base_price:
                price_diff = abs(latest_price - base_price) / base_price
            else:
                price_diff = 1.0   # 기준가 없으면 거리순

            apt_candidates.append({
                "name":        nm,
                "lat":         lat2,
                "lon":         lon2,
                "dist":        dist,
                "trades":      trades,
                "price_diff":  price_diff,
                "marker_type": "apt",
            })

        # 가격 ±40% 이내 우선 → 나머지 거리순
        in_range  = sorted([c for c in apt_candidates if c["price_diff"] <= 0.4],
                           key=lambda x: x["price_diff"])
        out_range = sorted([c for c in apt_candidates if c["price_diff"] > 0.4],
                           key=lambda x: x["dist"])
        fill_list = (in_range + out_range)[:needed]

        for apt in fill_list:
            nm = apt["name"]
            result[nm] = {
                "trades":      apt["trades"],
                "dist":        apt["dist"],
                "lat":         apt["lat"],
                "lon":         apt["lon"],
                "jibun":       "",
                "umd_nm":      "",
                "marker_type": "apt",
            }
            print(f"  ✅ [아파트 보완] {nm}: ({apt['dist']:.0f}m, 가격차 {apt['price_diff']*100:.0f}%)")

    print(f"  최종 주변 마커: {len(result)}개 "
          f"(오피스텔 {sum(1 for v in result.values() if v['marker_type']=='offi')}개 "
          f"+ 아파트 {sum(1 for v in result.values() if v['marker_type']=='apt')}개)")
    return result


def generate_images(detail, sales, nearby_sales, nearby_jeonse, output_dir, auto_mode=False, supply_map=None, jeonse=None, property_type="apt"):
    """이미지 4장 생성 (Selenium + 로컬 HTTP 서버)"""
    import http.server
    import socketserver

    lat, lon  = detail["lat"], detail["lon"]
    apt_name  = detail["kaptName"]
    dong      = detail.get("dong", "")
    os.makedirs(output_dir, exist_ok=True)

    # SEO 파일명 생성 — 하이픈 구분, 키워드 포함
    def _seo_name(*parts):
        """파트들을 하이픈으로 연결, 공백/특수문자 제거"""
        joined = "-".join(p.strip() for p in parts if p and p.strip())
        return re.sub(r"[\\/:*?\"<>|]", "", joined)  # 윈도우 금지 문자 제거

    prefix   = f"{dong}-" if dong else ""
    seo_base = _seo_name(dong, apt_name) if dong else apt_name

    # 주변 단지별 supply_map — 각 단지의 도로명으로 개별 조회
    if supply_map:
        for info in nearby_sales.values():
            if "supply_map" not in info:
                info["supply_map"] = supply_map   # 기본값: 기준 단지 것
        for info in nearby_jeonse.values():
            if "supply_map" not in info:
                info["supply_map"] = supply_map

    # 주변 단지별 정확한 공급면적 개별 조회 (DB → 건축물대장 API)
    for nm, info in list(nearby_sales.items()):
        sm = _get_supply_map_for_nearby(nm)
        if sm:
            info["supply_map"] = sm
    for nm, info in list(nearby_jeonse.items()):
        sm = _get_supply_map_for_nearby(nm)
        if sm:
            info["supply_map"] = sm

    # 로컬 HTTP 서버 시작 (output_dir 기준)
    PORT = 18765
    class Handler(http.server.SimpleHTTPRequestHandler):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, directory=output_dir, **kwargs)
        def log_message(self, *args):
            pass

    # 포트 재사용 허용 — 이전 실행이 포트를 점유하고 있어도 정상 시작
    socketserver.TCPServer.allow_reuse_address = True
    server = socketserver.TCPServer(("", PORT), Handler)
    thread = threading.Thread(target=server.serve_forever)
    thread.daemon = True
    thread.start()
    print(f"  로컬 서버 시작 (포트 {PORT})")

    type_word = "오피스텔" if property_type == "offi" else "아파트"

    print("  브라우저 시작...")
    driver = None
    paths = {}
    try:
        driver = _get_driver()
        paths["simple_map"] = get_map_image(driver, lat, lon, apt_name, output_dir, PORT,
                                            seo_name=_seo_name(dong, apt_name, "위치지도"))
        paths["deal_label"] = get_label_map(driver, lat, lon, apt_name, nearby_sales, "deal", output_dir, PORT,
                                            seo_name=_seo_name(dong, apt_name, f"{type_word}-매매가-시세지도"))
        paths["rent_label"] = get_label_map(driver, lat, lon, apt_name, nearby_jeonse, "rent", output_dir, PORT,
                                            seo_name=_seo_name(dong, apt_name, f"{type_word}-전세가-시세지도"))
    except Exception as e:
        print(f"❌ 이미지 생성 중 오류: {e}")
    finally:
        if driver:
            try:
                driver.quit()
            except Exception as e:
                print(f"  ⚠️ 브라우저 종료 오류: {e}")
        try:
            server.shutdown()
        except Exception as e:
            print(f"  ⚠️ 서버 종료 오류: {e}")
        print("  브라우저 / 서버 종료")

    paths["chart"] = get_sales_chart(sales, apt_name, output_dir,
                                      jeonse=jeonse, supply_map=supply_map,
                                      seo_name=_seo_name(dong, apt_name, f"{type_word}-매매가-전세가-실거래가-추이"))
    return paths


# ========================================================
# docx 원고 생성 (기존 포맷 동일)
# ========================================================
def _begin_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(13)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return doc


def _add_section_title(doc, no, text):
    p = doc.add_paragraph()
    rno = p.add_run(f"{no}. ")
    rno.bold = True
    rno.font.size = Pt(16)
    rno.font.color.rgb = RGBColor(10, 132, 255)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(10)
    return p


def _add_sub_title(doc, text):
    """카드 구분 소제목 (최근 거래 / 직전 거래) — 텍스트 색상만"""
    is_recent = "최근" in text
    p = doc.add_paragraph()
    r = p.add_run(f"▶ {text}")
    r.bold = True
    r.font.size = Pt(13)
    # 최근 거래: 그린 / 직전 거래: 슬레이트 블루
    r.font.color.rgb = RGBColor(15, 110, 86) if is_recent else RGBColor(51, 102, 153)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _add_bullet(doc, text, bold_text=None):
    p = doc.add_paragraph()
    dot = p.add_run("● ")
    dot.font.size = Pt(11)
    dot.font.color.rgb = RGBColor(128, 0, 32)
    if bold_text and bold_text in text:
        before, after = text.split(bold_text, 1)
        p.add_run(before)
        b = p.add_run(bold_text)
        b.bold = True
        p.add_run(after)
    else:
        p.add_run(text)
    p.paragraph_format.line_spacing = 1.15
    return p


def _add_card(doc, row, apt_name, card_type="trade", supply_map=None):
    """
    통합 카드 렌더링 — 매매/전세/월세 공통
    card_type: "trade" | "rent" | "monthly"
    """
    para  = doc.add_paragraph()
    exclu = row.get("excluUseAr", "")

    # 면적
    try:
        exclu_f   = float(str(exclu).replace(",", "").strip())
        exclu_str = f"{exclu_f:.2f}㎡"
    except (ValueError, TypeError):
        exclu_str = str(exclu) + "㎡" if exclu else "정보없음"
    supply_val = get_supply_from_map(exclu, supply_map) if supply_map else None
    area_str   = f"전용 {exclu_str} / 공급 {supply_val:.2f}㎡" if supply_val else f"전용 {exclu_str}"

    # 날짜
    date = (f"{row.get('dealYear','')}"
            f".{str(row.get('dealMonth','')).zfill(2)}"
            f".{str(row.get('dealDay','')).zfill(2)}")
    floor = row.get("floor", "?")

    # 가격 필드 — card_type에 따라 분기 (통일 헬퍼 사용)
    if card_type == "trade":
        price_label = "거래금액: "
        price_val   = format_price_word(row.get("dealAmount"))
        gbn         = row.get("dealingGbn", "")
        extra_items = [("거래유형: ", gbn, True)] if gbn else []
    elif card_type == "rent":
        price_label = "전세금: "
        dep_raw, _ = _resolve_deposit(row)
        price_val   = format_price_word(dep_raw)
        extra_items = []
    else:  # monthly
        price_label = "보증금/월세: "
        dep_raw, mon_raw = _resolve_monthly(row)
        dep_val = format_price_word(dep_raw)
        mon_val = format_price_word(mon_raw)
        price_val   = f"{dep_val}/{mon_val}"
        extra_items = []

    items = [
        ("단지명: ",   apt_name,        True),
        ("면적: ",     area_str,        True),
        ("거래일: ",   date,            False),
        (price_label, price_val,       True),
        ("층: ",       f"{floor}층",   False),
    ] + extra_items

    for i, (label, val, bold) in enumerate(items):
        d = para.add_run("● ")
        d.font.size = Pt(11)
        d.font.color.rgb = RGBColor(128, 0, 32)
        para.add_run(label).font.size = Pt(13)
        vr = para.add_run(val)
        vr.font.size = Pt(13)
        vr.bold = bold
        if i < len(items) - 1:
            para.add_run("\n")
    para.paragraph_format.line_spacing = 1.15


# 하위 호환 래퍼 (기존 호출부 변경 없이 유지)
def _add_trade_card(doc, row, apt_name, supply_map=None):
    _add_card(doc, row, apt_name, card_type="trade", supply_map=supply_map)

def _add_rent_card(doc, row, apt_name, supply_map=None):
    _add_card(doc, row, apt_name, card_type="rent", supply_map=supply_map)

def _add_monthly_card(doc, row, apt_name, supply_map=None):
    _add_card(doc, row, apt_name, card_type="monthly", supply_map=supply_map)


def _generate_tags(kapt_name, dong, gu, sales, jeonse, property_type="apt"):
    is_offi   = (property_type == "offi")
    type_word = "오피스텔" if is_offi else "아파트"
    tags = [f"#{kapt_name}", f"#{kapt_name}{type_word}", f"#{kapt_name}실거래가"]
    if dong:
        tags += [f"#{dong}", f"#{dong}{type_word}", f"#{dong}실거래가"]
    if gu:
        tags += [f"#{gu}", f"#{gu}{type_word}매매", f"#{gu}전세"]
    if is_offi:
        tags += ["#오피스텔매매", "#오피스텔전세", "#오피스텔실거래가", "#국토부실거래가"]
    else:
        tags += ["#아파트매매", "#아파트전세", "#실거래가", "#국토부실거래가"]
    return tags


def _validate_data(apt_name, sales, jeonse, wolse, supply_map, property_type):
    """
    원고 생성 전 데이터 정합성 검증
    - 올바른 필드에서 값을 가져왔는지 확인
    - 이상값(0원, 0㎡, 필드 누락) 감지
    - 결과는 로그로만 출력, 원고 생성은 항상 진행
    """
    ok   = "  ✅"
    warn = "  ⚠️"
    err  = "  ❌"
    issues = []

    print(f"\n{'─'*50}")
    print(f"[검증] {apt_name}")

    # ── 매매 검증 ─────────────────────────────────────
    if not sales:
        issues.append(f"{err} 매매 데이터 없음")
    else:
        # 최근 3건 기준
        recent = sorted(sales,
            key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2),
                           str(x.get("dealDay","")).zfill(2)), reverse=True)[:3]
        for i, row in enumerate(recent):
            # dealAmount 필드 검증
            amt_raw = row.get("dealAmount", "")
            amt     = parse_price(amt_raw)
            exclu   = row.get("excluUseAr", "")
            date    = f"{row.get('dealYear','')}.{str(row.get('dealMonth','')).zfill(2)}"

            if not amt or amt <= 0:
                issues.append(f"{err} 매매 {i+1}번: dealAmount 이상값 ({amt_raw!r})")
            elif amt < 500:
                issues.append(f"{err} 매매 {i+1}번: 가격이 너무 낮음 ({amt}만원) — 단위 오류 의심")
            else:
                # 공급면적 매핑 확인
                supply_val = get_supply_from_map(exclu, supply_map) if supply_map else None
                area_disp  = f"공급{supply_val:.0f}㎡" if supply_val else f"전용{exclu}㎡"
                print(f"{ok} 매매 [{date}] {area_disp} → {format_price_word(amt_raw)}"
                      f"{'  (공급면적 매핑 없음)' if not supply_val else ''}")

            if not exclu:
                issues.append(f"{warn} 매매 {i+1}번: excluUseAr(전용면적) 없음")

    # ── 전세 검증 ─────────────────────────────────────
    if not jeonse:
        print(f"{warn} 전세 데이터 없음")
    else:
        recent_j = sorted(jeonse,
            key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2),
                           str(x.get("dealDay","")).zfill(2)), reverse=True)[:3]
        for i, row in enumerate(recent_j):
            # 통일 헬퍼로 보증금 필드 결정
            val_raw, used_field = _resolve_deposit(row)

            val    = parse_price(val_raw)
            exclu  = row.get("excluUseAr", "")
            date   = f"{row.get('dealYear','')}.{str(row.get('dealMonth','')).zfill(2)}"

            if not val or val <= 0:
                issues.append(f"{err} 전세 {i+1}번: {used_field} 이상값 ({val_raw!r})")
            elif val < 500:
                issues.append(f"{err} 전세 {i+1}번: 보증금이 너무 낮음 ({val}만원)")
            else:
                supply_val = get_supply_from_map(exclu, supply_map) if supply_map else None
                area_disp  = f"공급{supply_val:.0f}㎡" if supply_val else f"전용{exclu}㎡"
                field_note = "" if used_field == "deposit" else "  (preDeposit 사용)"
                print(f"{ok} 전세 [{date}] {area_disp} → {format_price_word(val_raw)}{field_note}")

    # ── 공급면적 검증 ──────────────────────────────────
    if not supply_map:
        issues.append(f"{warn} 공급면적 데이터 없음 — 전용면적으로만 표시됨")
    else:
        print(f"{ok} 공급면적 {len(supply_map)}개 평형 확인")

    # ── 매매/전세 가격 비율 검증 ───────────────────────
    if sales and jeonse:
        try:
            last_sale = sorted(sales,
                key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2)),
                reverse=True)[0]
            last_jeon = sorted(jeonse,
                key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2)),
                reverse=True)[0]
            sale_amt = parse_price(last_sale.get("dealAmount"))
            dep_raw  = last_jeon.get("deposit","")
            pre_raw  = last_jeon.get("preDeposit","")
            jeon_amt = parse_price(dep_raw if _is_valid_num(dep_raw) else pre_raw)
            if sale_amt and jeon_amt and sale_amt > 0:
                ratio = jeon_amt / sale_amt * 100
                if ratio > 100:
                    issues.append(f"{err} 전세가({jeon_amt:,}만)가 매매가({sale_amt:,}만)보다 높음 — 필드 오류 가능성")
                elif ratio > 95:
                    issues.append(f"{warn} 전세가율 {ratio:.0f}% — 매우 높음 (실제 여부 확인 권장)")
                else:
                    print(f"{ok} 전세가율 {ratio:.0f}% — 정상 범위")
        except Exception:
            pass

    # ── 결과 출력 ──────────────────────────────────────
    if issues:
        print(f"\n  발견된 이슈:")
        for issue in issues:
            print(f"  {issue}")
    else:
        print(f"{ok} 모든 항목 검증 통과")
    print(f"{'─'*50}\n")


def _line_emoji(line_name):
    """노선명 → 노선 색상 이모티콘"""
    LINE_COLOR = {
        # 수도권 지하철
        "1호선": "🔵", "2호선": "🟢", "3호선": "🟠", "4호선": "🔵",
        "5호선": "🟣", "6호선": "🟤", "7호선": "🟢", "8호선": "🩷",
        "9호선": "🟡",
        # 광역/특수
        "경의중앙선": "🩵", "경춘선": "🟢", "분당선": "🟡",
        "수인선": "🟡", "수인분당선": "🟡",
        "신분당선": "🔴", "공항철도": "🔵",
        "GTX-A": "⚫", "GTX-B": "⚫", "GTX-C": "⚫",
        "우이신설선": "🟢",
        # 부산
        "부산 도시철도 1호선": "🟠", "부산 도시철도 2호선": "🟠",
        "부산 도시철도 3호선": "🟤", "부산 도시철도 4호선": "🔵",
        # 대구
        "대구 도시철도 1호선": "🔴", "대구 도시철도 2호선": "🟢",
        "대구 도시철도 3호선": "🟠",
        # 인천
        "인천지하철 1호선": "🟢", "인천지하철 2호선": "🟡",
        # 코레일 기본
        "코레일": "🔵",
    }
    # 완전 일치 우선
    if line_name in LINE_COLOR:
        return LINE_COLOR[line_name]
    # 부분 매칭
    for key, emoji in LINE_COLOR.items():
        if key in line_name or line_name in key:
            return emoji
    return "🚇"  # 기본


def generate_documents(detail, sales, jeonse, wolse, nearby_sales, schools, output_dir, property_type="apt", supply_map=None, stations=None):
    """
    docx 원고 8개 생성
    01_블로그제목 / 02_아파트개요 / 03_매매카드 / 05_전세카드 /
    07_월세카드 / 08_학교정보 / 10_자료출처 / 11_연관태그
    property_type: "apt" | "offi"
    """
    os.makedirs(output_dir, exist_ok=True)
    is_offi    = (property_type == "offi")
    type_label = "오피스텔" if is_offi else "아파트"

    apt_name   = detail["kaptName"]
    dong       = detail.get("dong", "")
    gu         = detail.get("gu", "")
    use_date   = detail.get("kaptUsedate", "")
    households = str(int(float(detail.get("kaptdaCnt", "") or 0))) if detail.get("kaptdaCnt") else ""
    address    = detail.get("doroJuso", "")
    top_floor  = detail.get("kaptTopFloor", "")
    company    = detail.get("kaptBcompany", "")
    parking    = detail.get("kaptTotPkCnt", "")
    build_year = use_date[:4] if use_date else ""

    # 오피스텔: 준공년도를 실거래가 buildYear에서 보완
    if is_offi and not build_year and sales:
        build_year = sales[0].get("buildYear", "")[:4] if sales[0].get("buildYear") else ""

    print(f"\n원고 생성 중: {apt_name} ({type_label}) → {output_dir}")

    # ── 01 블로그 제목 ──────────────────────────────────
    if is_offi:
        apt_suffix = "" if apt_name.endswith("오피스텔") else " 오피스텔"
    else:
        apt_suffix = "" if apt_name.endswith("아파트") else " 아파트"
    prefix = f"{dong} " if dong else ""
    year   = datetime.now().year

    # 최근 매매가 / 전세가 추출 (제목용)
    def _price_short(data, field="dealAmount"):
        """최근 거래가를 'X.X억' 형태로 반환"""
        if not data:
            return None
        recent = sorted(data,
            key=lambda x: (x.get("dealYear",""), str(x.get("dealMonth","")).zfill(2)),
            reverse=True)
        for r in recent:
            if field == "deposit":
                dep = r.get("deposit","")
                pre = r.get("preDeposit","")
                val = parse_price(dep if _is_valid_num(dep) else pre)
            else:
                val = parse_price(r.get(field))
            if val and val > 0:
                eok = val / 10000
                if eok >= 1:
                    return f"{eok:.1f}억".rstrip("0억").rstrip(".") + "억" if eok != int(eok) else f"{int(eok)}억"
                return f"{val:,}만"
        return None

    sale_price = _price_short(sales, "dealAmount")
    jeon_price = _price_short(jeonse, "deposit") if jeonse else None

    # 가격 문자열 조합
    if sale_price and jeon_price:
        price_str = f"매매 {sale_price} · 전세 {jeon_price}"
    elif sale_price:
        price_str = f"매매 {sale_price}"
    else:
        price_str = "실거래가"

    name_full = f"{prefix}{apt_name}{apt_suffix}"
    # 단지명만 (suffix 없이) — 짧은 제목용
    name_short = f"{prefix}{apt_name}" if prefix else apt_name

    title_templates = [
        # 1. 질문형
        f"{apt_name}{apt_suffix} 지금 얼마? {price_str} [{year} 실거래가]",
        # 2. 정보형
        f"[{year} 최신] {gu if gu else ''} {name_full} {price_str} 실거래가 분석".strip(),
        # 3. 혼합형
        f"{name_full} {year}년 {price_str} 실거래가 시세 총정리",
    ]
    title = random.choice(title_templates)

    # 50자 초과 시 간소화 버전으로 교체
    # (네이버 모바일 기준 50자 이상이면 두 줄 가능)
    if len(title) > 50:
        short_templates = [
            f"{apt_name} {price_str} 실거래가 [{year}]",
            f"{name_short} {year} 매매·전세 실거래가",
            f"{apt_name} {year}년 실거래가 시세",
        ]
        title = random.choice(short_templates)

    doc = _begin_doc()
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(10, 132, 255)
    doc.save(os.path.join(output_dir, "01_블로그제목.docx"))

    # ── 02 개요 ────────────────────────────────────────
    doc = _begin_doc()
    _add_section_title(doc, 1, f"{prefix}{apt_name} {type_label} 개요")
    _add_bullet(doc, f"주소: {address}", address if address else None)
    _add_bullet(doc, f"준공년도: {build_year}년" if build_year else "준공년도: 정보없음")
    if not is_offi:
        # 아파트만 세대수/최고층/시공사 표시
        _add_bullet(doc, f"세대수: {households}세대" if households else "세대수: 정보없음",
                    f"{households}세대" if households else None)
        _add_bullet(doc, f"최고층: {top_floor}층" if top_floor else "최고층: 정보없음")
        _add_bullet(doc, f"시공사: {company}" if company else "시공사: 정보없음")

    # 역세권 정보 (가장 가까운 역 최대 2개)
    if stations:
        for s in stations[:2]:
            emoji = _line_emoji(s['line'])
            _add_bullet(doc,
                f"{emoji} 역세권: {s['name']} ({s['line']}) 도보 {s['walk_min']}분 ({s['dist']:.0f}m)",
                f"도보 {s['walk_min']}분")
    doc.save(os.path.join(output_dir, "02_아파트개요.docx"))

    # ── 03 매매 카드 ────────────────────────────────────
    doc = _begin_doc()
    _add_section_title(doc, 2, f"{prefix}{apt_name} 최근 {type_label} 매매 실거래가")
    if sales:
        # 3년 내 최고가 (주력 평형 기준)
        all_prices = [parse_price(r.get("dealAmount")) for r in sales if parse_price(r.get("dealAmount"))]
        if all_prices:
            max_price = max(all_prices)
            max_row   = next(r for r in sales if parse_price(r.get("dealAmount")) == max_price)
            max_date  = f"{max_row.get('dealYear','')}.{str(max_row.get('dealMonth','')).zfill(2)}"
            max_exclu = max_row.get("excluUseAr","")
            max_supply = get_supply_from_map(max_exclu, supply_map) if supply_map else None
            area_disp  = f"공급 {max_supply:.0f}㎡" if max_supply else f"전용 {max_exclu}㎡"
            p = doc.add_paragraph()
            r1 = p.add_run("📌 3년 내 최고가  ")
            r1.bold = True
            r1.font.size = Pt(12)
            r1.font.color.rgb = RGBColor(200, 50, 50)
            r2 = p.add_run(f"{format_price_word(max_price)}  ({area_disp} / {max_date})")
            r2.bold = True
            r2.font.size = Pt(13)
            r2.font.color.rgb = RGBColor(200, 50, 50)
            p.paragraph_format.space_after = Pt(8)
            doc.add_paragraph()

        top = sorted(sales, key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2),
                                           str(x.get("dealDay","")).zfill(2)), reverse=True)
        labels = ["최근 거래", "직전 거래"]
        for i, row in enumerate(top[:2]):
            _add_sub_title(doc, labels[i])
            _add_trade_card(doc, row, apt_name, supply_map=supply_map)
            doc.add_paragraph()
    else:
        _add_bullet(doc, "※ 매매 데이터 없음")
    doc.save(os.path.join(output_dir, "03_매매카드.docx"))

    # ── 05 전세 카드 ────────────────────────────────────
    doc = _begin_doc()
    _add_section_title(doc, 4, f"{prefix}{apt_name} 최근 {type_label} 전세 실거래가")
    if jeonse:
        top = sorted(jeonse, key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2),
                                            str(x.get("dealDay","")).zfill(2)), reverse=True)
        labels = ["최근 거래", "직전 거래"]
        for i, row in enumerate(top[:2]):
            _add_sub_title(doc, labels[i])
            _add_rent_card(doc, row, apt_name, supply_map=supply_map)
            doc.add_paragraph()
    else:
        _add_bullet(doc, "※ 전세 데이터 없음")
    doc.save(os.path.join(output_dir, "05_전세카드.docx"))

    # ── 07 월세 카드 ────────────────────────────────────
    doc = _begin_doc()
    _add_section_title(doc, 6, f"{prefix}{apt_name} 최근 {type_label} 월세 실거래가")
    if wolse:
        top = sorted(wolse, key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2),
                                           str(x.get("dealDay","")).zfill(2)), reverse=True)
        labels = ["최근 거래", "직전 거래"]
        for i, row in enumerate(top[:2]):
            _add_sub_title(doc, labels[i])
            _add_monthly_card(doc, row, apt_name, supply_map=supply_map)
            doc.add_paragraph()
    else:
        _add_bullet(doc, "※ 월세 데이터 없음")
    doc.save(os.path.join(output_dir, "07_월세카드.docx"))

    # ── 04 매매 전문가 의견 ─────────────────────────────
    doc = _begin_doc()
    _add_section_title(doc, 3, f"{prefix}{apt_name} 매매 전문가 의견")

    # 휙 브랜드 헤더
    p = doc.add_paragraph()
    r = p.add_run("🔷 휙(Hwik) 전문가 의견")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(10, 132, 255)
    p.paragraph_format.space_after = Pt(6)

    if sales:
        top_sale = sorted(sales,
            key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2),
                           str(x.get("dealDay","")).zfill(2)), reverse=True)
        recent = top_sale[0]
        exclu  = recent.get("excluUseAr","")
        sup    = get_supply_from_map(exclu, supply_map) if supply_map else None
        area   = f"{sup:.0f}㎡(공급)" if sup else f"{exclu}㎡(전용)"
        date   = f"{recent.get('dealYear','')}.{str(recent.get('dealMonth','')).zfill(2)}"
        price  = format_price_word(recent.get("dealAmount"))

        # 주변 단지 최근 매매 2곳 먼저 수집 (DB에 있는 단지만)
        nearby_list = []
        for nm, info in nearby_sales.items():
            trades = info.get("trades", [])
            if not trades:
                continue
            t = sorted(trades,
                key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2)),
                reverse=True)[0]
            amt = parse_price(t.get("dealAmount"))
            if not amt:
                continue
            ex2 = t.get("excluUseAr","")
            nearby_supply = _get_supply_map_for_nearby(nm)
            if not nearby_supply:
                continue  # DB에 없는 단지 → 면적 불확실 → 스킵
            s2 = get_supply_from_map(ex2, nearby_supply)
            if not s2:
                continue  # 공급면적 매핑 실패 → 스킵
            a2 = f"{s2:.0f}㎡(공급)"
            d2 = f"{t.get('dealYear','')}.{str(t.get('dealMonth','')).zfill(2)}"
            nearby_list.append((nm, a2, format_price_word(amt), d2))
            if len(nearby_list) >= 2:
                break
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{apt_name} {area}의 최근 매매 실거래가는 ").font.size = Pt(12)
        pr = p.add_run(price); pr.bold = True; pr.font.size = Pt(12)
        p.add_run(f"({date})").font.size = Pt(12)
        p.add_run("입니다.").font.size = Pt(12)
        if nearby_list:
            p.add_run(" 인근 ").font.size = Pt(12)
            for i, (nm2, a2, pr2, d2) in enumerate(nearby_list):
                if i > 0:
                    p.add_run(", ").font.size = Pt(12)
                p.add_run(f"{nm2} {a2} ").font.size = Pt(12)
                rb = p.add_run(pr2); rb.bold = True; rb.font.size = Pt(12)
                p.add_run(f"({d2})").font.size = Pt(12)
            p.add_run("에 거래됐습니다.").font.size = Pt(12)
    else:
        doc.add_paragraph("※ 매매 데이터 없음")
    doc.save(os.path.join(output_dir, "04_매매전문가의견.docx"))

    # ── 06 전세 전문가 의견 ─────────────────────────────
    doc = _begin_doc()
    _add_section_title(doc, 5, f"{prefix}{apt_name} 전세 전문가 의견")

    p = doc.add_paragraph()
    r = p.add_run("🔷 휙(Hwik) 전문가 의견")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(10, 132, 255)
    p.paragraph_format.space_after = Pt(6)

    if jeonse:
        top_jeon = sorted(jeonse,
            key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2),
                           str(x.get("dealDay","")).zfill(2)), reverse=True)
        recent_j = top_jeon[0]
        exclu_j  = recent_j.get("excluUseAr","")
        sup_j    = get_supply_from_map(exclu_j, supply_map) if supply_map else None
        area_j   = f"{sup_j:.0f}㎡(공급)" if sup_j else f"{exclu_j}㎡(전용)"
        date_j   = f"{recent_j.get('dealYear','')}.{str(recent_j.get('dealMonth','')).zfill(2)}"
        dep_raw  = recent_j.get("deposit","")
        pre_raw  = recent_j.get("preDeposit","")
        price_j  = format_price_word(dep_raw if _is_valid_num(dep_raw) else pre_raw)

        # 전세가율 계산 — 같은 평형 매매가 기준
        ratio_str = ""
        if sales:
            jeon_exclu = recent_j.get("excluUseAr", "")
            try:
                jeon_exclu_f = float(str(jeon_exclu).replace(",","").strip()) if jeon_exclu else 0
            except:
                jeon_exclu_f = 0
            # 같은 평형(2㎡ 이내) 매매 중 가장 최근
            same_pyeong_sales = []
            if jeon_exclu_f > 0:
                same_pyeong_sales = [
                    r for r in sales
                    if abs(float(str(r.get("excluUseAr","0")).replace(",","").strip() or 0) - jeon_exclu_f) <= 2.0
                ]
            if not same_pyeong_sales:
                same_pyeong_sales = sales  # 같은 평형 없으면 전체 폴백
            sale_rec = sorted(same_pyeong_sales,
                key=lambda x: (x.get("dealYear",""), x.get("dealMonth","")),
                reverse=True)[0]
            sale_amt = parse_price(sale_rec.get("dealAmount"))
            jeon_amt = parse_price(dep_raw if _is_valid_num(dep_raw) else pre_raw)
            if sale_amt and jeon_amt and sale_amt > 0:
                ratio = jeon_amt / sale_amt * 100
                if ratio <= 100:  # 100% 초과는 표시 안 함 (데이터 오류 가능성)
                    ratio_str = f"{ratio:.0f}%"
                else:
                    print(f"  ⚠️ 전세가율 {ratio:.0f}% 이상값 → 표시 생략 (평형 불일치 가능)")

        # 주변 전세 2곳 수집 (DB에 있는 단지만)
        nearby_jeon_list = []
        for nm, info in nearby_sales.items():
            trades = info.get("trades", [])
            jeon_t = [t for t in trades if not t.get("monthlyRent") or t.get("monthlyRent") == "0"]
            if not jeon_t:
                continue
            t = sorted(jeon_t,
                key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2)),
                reverse=True)[0]
            dep = t.get("deposit","")
            pre = t.get("preDeposit","")
            amt = parse_price(dep if _is_valid_num(dep) else pre)
            if not amt:
                continue
            ex2 = t.get("excluUseAr","")
            nearby_supply2 = _get_supply_map_for_nearby(nm)
            if not nearby_supply2:
                continue  # DB에 없는 단지 → 스킵
            s2 = get_supply_from_map(ex2, nearby_supply2)
            if not s2:
                continue  # 공급면적 매핑 실패 → 스킵
            a2 = f"{s2:.0f}㎡(공급)"
            d2 = f"{t.get('dealYear','')}.{str(t.get('dealMonth','')).zfill(2)}"
            nearby_jeon_list.append((nm, a2, format_price_word(amt), d2))
            if len(nearby_jeon_list) >= 2:
                break
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{apt_name} {area_j}의 최근 전세 실거래가는 ").font.size = Pt(12)
        pr = p.add_run(price_j); pr.bold = True; pr.font.size = Pt(12)
        p.add_run(f"({date_j})").font.size = Pt(12)
        if ratio_str:
            p.add_run(f"이며 전세가율은 ").font.size = Pt(12)
            rb2 = p.add_run(ratio_str)
            rb2.bold = True; rb2.font.size = Pt(12)
        p.add_run("입니다.").font.size = Pt(12)
        if nearby_jeon_list:
            p.add_run(" 인근 ").font.size = Pt(12)
            for i, (nm2, a2, pr2, d2) in enumerate(nearby_jeon_list):
                if i > 0:
                    p.add_run(", ").font.size = Pt(12)
                p.add_run(f"{nm2} {a2} ").font.size = Pt(12)
                rb = p.add_run(pr2); rb.bold = True; rb.font.size = Pt(12)
                p.add_run(f"({d2})").font.size = Pt(12)
            p.add_run("에 전세 계약됐습니다.").font.size = Pt(12)
    else:
        doc.add_paragraph("※ 전세 데이터 없음")
    doc.save(os.path.join(output_dir, "06_전세전문가의견.docx"))

    # ── 09 종합 의견 ────────────────────────────────────
    doc = _begin_doc()
    _add_section_title(doc, 8, f"{prefix}{apt_name} 종합 의견")

    p = doc.add_paragraph()
    r = p.add_run("🔷 휙(Hwik) 전문가 의견")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(10, 132, 255)
    p.paragraph_format.space_after = Pt(6)

    # 매매 + 전세 + 3년변동 → 하나의 단락으로
    if sales or jeonse:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

        # 매매가
        if sales:
            rec_s  = sorted(sales,
                key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2)),
                reverse=True)[0]
            ex_s   = rec_s.get("excluUseAr","")
            sup_s  = get_supply_from_map(ex_s, supply_map) if supply_map else None
            area_s = f"{sup_s:.0f}㎡(공급)" if sup_s else f"{ex_s}㎡(전용)"
            dt_s   = f"{rec_s.get('dealYear','')}.{str(rec_s.get('dealMonth','')).zfill(2)}"
            pr_s   = format_price_word(rec_s.get("dealAmount"))
            p.add_run(f"{apt_name} {area_s} 매매 실거래가는 ").font.size = Pt(12)
            r2 = p.add_run(pr_s); r2.bold = True; r2.font.size = Pt(12)
            p.add_run(f"({dt_s})").font.size = Pt(12)
            p.add_run("입니다. ").font.size = Pt(12)

        # 전세 + 전세가율
        if jeonse:
            rec_j  = sorted(jeonse,
                key=lambda x: (x.get("dealYear",""), x.get("dealMonth","").zfill(2)),
                reverse=True)[0]
            ex_j   = rec_j.get("excluUseAr","")
            sup_j2 = get_supply_from_map(ex_j, supply_map) if supply_map else None
            area_j2= f"{sup_j2:.0f}㎡(공급)" if sup_j2 else f"{ex_j}㎡(전용)"
            dt_j   = f"{rec_j.get('dealYear','')}.{str(rec_j.get('dealMonth','')).zfill(2)}"
            dep_r  = rec_j.get("deposit","")
            pre_r  = rec_j.get("preDeposit","")
            pr_j   = format_price_word(dep_r if _is_valid_num(dep_r) else pre_r)
            p.add_run(f"전세 실거래가는 {area_j2} 기준 ").font.size = Pt(12)
            r3 = p.add_run(pr_j); r3.bold = True; r3.font.size = Pt(12)
            p.add_run(f"({dt_j})").font.size = Pt(12)
            if sales:
                try:
                    j_exclu_f = float(str(ex_j).replace(",","").strip()) if ex_j else 0
                except:
                    j_exclu_f = 0
                same_py = [r for r in sales
                           if j_exclu_f > 0 and
                           abs(float(str(r.get("excluUseAr","0")).replace(",","").strip() or 0) - j_exclu_f) <= 2.0]
                if not same_py:
                    same_py = sales
                sale_v = parse_price(sorted(same_py,
                    key=lambda x: (x.get("dealYear",""), x.get("dealMonth","")),
                    reverse=True)[0].get("dealAmount"))
                jeon_v = parse_price(dep_r if _is_valid_num(dep_r) else pre_r)
                if sale_v and jeon_v and sale_v > 0:
                    ratio = jeon_v / sale_v * 100
                    if ratio <= 100:
                        p.add_run(f"으로 전세가율은 ").font.size = Pt(12)
                        r4 = p.add_run(f"{ratio:.0f}%")
                        r4.bold = True; r4.font.size = Pt(12)
            p.add_run("입니다. ").font.size = Pt(12)

        # 3년 변동
        if sales and len(sales) >= 2:
            sorted_s = sorted(sales,
                key=lambda x: (x.get("dealYear",""), x.get("dealMonth","")))
            oldest = next((parse_price(r.get("dealAmount")) for r in sorted_s
                           if parse_price(r.get("dealAmount"))), None)
            newest = next((parse_price(r.get("dealAmount")) for r in reversed(sorted_s)
                           if parse_price(r.get("dealAmount"))), None)
            if oldest and newest:
                diff = newest - oldest
                sign = "+" if diff >= 0 else ""
                old_date = f"{sorted_s[0].get('dealYear','')}.{str(sorted_s[0].get('dealMonth','')).zfill(2)}"
                new_date = f"{sorted_s[-1].get('dealYear','')}.{str(sorted_s[-1].get('dealMonth','')).zfill(2)}"
                p.add_run(f"3년간({old_date} → {new_date}) 매매가 변동은 ").font.size = Pt(12)
                r5 = p.add_run(f"{sign}{format_price_word(abs(diff))}")
                r5.bold = True; r5.font.size = Pt(12)
                p.add_run("입니다.").font.size = Pt(12)
    else:
        doc.add_paragraph("※ 데이터 없음")
    doc.save(os.path.join(output_dir, "09_종합의견.docx"))
    doc = _begin_doc()
    _add_section_title(doc, 7, f"{prefix}{apt_name} 주변 학교 정보")
    if schools:
        # 초중고 각 1개씩 (가장 가까운 것)
        shown = {}
        for s in schools:  # 이미 거리순 정렬됨
            stype = s["type"]
            if stype in ("초등학교", "중학교", "고등학교") and stype not in shown:
                shown[stype] = s
            if len(shown) == 3:
                break
        # 초→중→고 순서로 출력
        for stype in ["초등학교", "중학교", "고등학교"]:
            if stype in shown:
                s = shown[stype]
                walk_min = max(1, round(s["dist"] / 80))
                _add_bullet(doc, f"{stype}: {s['name']} (도보 {walk_min}분, {s['dist']:.0f}m)")
        if not shown:
            _add_bullet(doc, "※ 반경 3km 내 학교 정보 없음")
    else:
        _add_bullet(doc, "※ 반경 3km 내 학교 정보 없음")
    doc.save(os.path.join(output_dir, "08_학교정보.docx"))

    # ── 10 자료출처 ─────────────────────────────────────
    doc = _begin_doc()
    _add_section_title(doc, 9, "자료출처")
    now = datetime.now()
    p = doc.add_paragraph()
    r = p.add_run(f"국토교통부 실거래가 (최근 3년, {now.year}년 {now.month}월까지)")
    r.bold = True
    r.font.size = Pt(13)
    p.add_run(", 카카오맵 지도")

    engagement = random.choice([
        "\n✨ 정보가 도움이 되셨다면 공감(💗) 한번 눌러주세요!\n궁금하신 점이나 관심지역은 댓글로 남겨주세요 😊",
        "\n─────────────────\n💗 유익했다면 공감 클릭!\n💬 궁금한 단지는 댓글로!\n다음 포스팅 주제 선정에 반영됩니다 ✨",
        "\n📊 매주 새로운 아파트 분석이 업데이트됩니다\n\n💗 이 정보가 유익하셨다면 공감 버튼 클릭!\n💬 알고 싶은 단지를 댓글로 남겨주세요",
    ])
    doc.add_paragraph(engagement)
    doc.save(os.path.join(output_dir, "10_자료출처.docx"))

    # ── 11 연관태그 ─────────────────────────────────────
    doc = _begin_doc()
    tags = _generate_tags(apt_name, dong, gu, sales, jeonse, property_type)
    p = doc.add_paragraph()
    r = p.add_run(" ".join(tags))
    r.bold = True
    r.font.size = Pt(14)
    doc.save(os.path.join(output_dir, "11_연관태그.docx"))

    print(f"✅ 워드 문서 11개 저장 완료 → {output_dir}")


# ========================================================
# 전체 파이프라인
# ========================================================
def run_pipeline(user_input, output_base=None, auto_mode=False, photo_paths=None):
    if output_base is None:
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        output_base = os.path.join(desktop, "원고")
    print(f"\n{'='*50}")
    print(f"입력: {user_input}")
    print(f"{'='*50}\n")

    # 1. 파싱
    parsed = parse_input_with_claude(user_input)
    if not parsed:
        print("❌ 파싱 실패")
        return

    # 2. 단지 조회 — DB 우선, 실패 시 기존 API 폴백
    forced_offi = False
    property_hint = parsed.get("property_hint", "apt")

    # ── 2-A. Supabase DB 검색 (1순위) ──────────────────────
    detail = find_apt_from_db(
        parsed.get("apt_name", ""),
        gu=parsed.get("gu", ""),
        dong=parsed.get("dong", ""),
    )

    if detail:
        # DB에서 property_type 확정
        db_prop_type = detail.pop("_db_prop_type", "apt")
        if db_prop_type == "offi" or property_hint == "officetel":
            forced_offi = True
            # 오피스텔인데 doroJuso 없으면 카카오로 보완
            if not detail.get("doroJuso"):
                print("  [오피스텔] 도로명 주소 없음 → 카카오로 보완")
                doc = search_kakao_keyword(
                    f"{parsed.get('gu','')} {parsed.get('dong','')} {detail['kaptName']}",
                    allow_offi=True
                )
                if doc:
                    detail["doroJuso"] = (
                        doc.get("road_address_name") or doc.get("address_name", "")
                    )
    else:
        # ── 2-B. 기존 API 폴백 ──────────────────────────────
        print("  → DB 미발견 — 기존 API 경로로 조회")

        if property_hint == "officetel":
            print("  → 오피스텔 입력 감지 — 오피스텔 경로 직행")
            detail = find_apt_offi(parsed)
            forced_offi = True
            if not detail:
                print("❌ 단지 조회 완전 실패 (오피스텔 경로)")
                return
        else:
            detail = find_apt(parsed)

            if detail:
                input_clean  = parsed.get("apt_name", "").replace(" ", "")
                result_clean = detail["kaptName"].replace(" ", "")
                name_match = ((input_clean in result_clean) or (result_clean in input_clean)
                              or _fuzzy_score(input_clean, result_clean) >= 0.72)
                if not name_match:
                    print(f"⚠️  단지명 불일치: 입력='{input_clean}' / 조회='{result_clean}'")
                    print("   → 오피스텔 경로로 전환...")
                    detail = None

            if not detail:
                print("아파트 단지 조회 실패 → 오피스텔 경로 시도...")
                detail = find_apt_offi(parsed)
                forced_offi = True
                if not detail:
                    print("❌ 단지 조회 완전 실패 (아파트/오피스텔 모두 불가)")
                    return

    detail["dong"] = parsed.get("dong", "")
    detail["gu"]   = parsed.get("gu", "")

    apt_name = detail["kaptName"]
    lawd_cd  = detail["bjd_code"][:5]
    road_address = detail.get("doroJuso", "")

    # 도로명+본번 추출 (매매 PRIMARY 필터용)
    road_nm, road_bonbun = _extract_road_info(road_address)

    # 지번 추출 (전월세 PRIMARY 필터용 — 전월세 API에 roadNm 없음)
    jibun_str = ""
    umd_nm_str = ""
    if road_address:
        _, _, _, jibun_raw, umd_nm_raw = address_to_info(road_address)
        jibun_str  = jibun_raw  or ""
        umd_nm_str = umd_nm_raw or ""
        if jibun_str:
            print(f"  지번: {umd_nm_str} {jibun_str}")

    # ── danji_pages 캐시 시도 (API 호출 대폭 감소) ──
    danji_cache = fetch_danji_page_data(apt_name, gu=detail.get("gu",""), dong=detail.get("dong",""))

    if danji_cache and danji_cache.get("price_history"):
        print(f"\n⚡ danji_pages 캐시 사용 — API 호출 스킵")
        sales = _danji_page_to_sales(danji_cache)
        property_type = "offi" if forced_offi else "apt"
        if sales:
            type_label = "아파트" if property_type == "apt" else "오피스텔"
            print(f"  🏢 건물 유형: {type_label} (매매 {len(sales)}건)")
            jeonse, wolse = _danji_page_to_rent(danji_cache)
            print(f"  전세 {len(jeonse)}건, 월세 {len(wolse)}건")
            schools = danji_cache.get("nearby_school") or []
            print(f"  학교 {len(schools)}개")
            stations_nearby = danji_cache.get("nearby_subway") or []
            print(f"  지하철 {len(stations_nearby)}개")
            nearby_sales = _danji_page_to_nearby(danji_cache)
            nearby_jeonse = {}
            print(f"  주변 단지 {len(nearby_sales)}개")
        else:
            print("  ⚠️ 캐시에 매매 데이터 없음 → API 폴백")
            danji_cache = None

    if not danji_cache or not danji_cache.get("price_history") or not sales:
        print(f"\n📡 API 직접 조회 (캐시 없음)")

        # 3. 매매 실거래가
        # 아파트: 도로명 PRIMARY → jibun fallback / 오피스텔: 키워드 PRIMARY
        sales, property_type = fetch_sales_auto(apt_name, lawd_cd, months=36,
                                            force_offi=forced_offi,
                                            road_address=road_address,
                                            jibun=jibun_str, umd_nm=umd_nm_str)

    if property_type is None:
        print("❌ 실거래가 데이터 없음")
        print("   아파트/오피스텔 모두 해당 없음 (연립·다세대·빌라·단독주택은 지원하지 않습니다)")
        return

    # 실거래가 API에서 반환된 실제 단지명으로 보정 (오피스텔 경로에서 중요)
    # 예: 입력 "효성서너스빌에코" → API 응답 "효성써너스빌에코오피스텔" → 정식 명칭 사용
    # 오피스텔 경로에서만 단지명 보정 (아파트는 국토부 DB 이름이 정확)
    if sales and forced_offi:
        name_field = "aptNm"
        name_counts = Counter(
            row.get(name_field, "").strip()
            for row in sales if row.get(name_field, "").strip()
        )
        real_name = name_counts.most_common(1)[0][0] if name_counts else ""
        if real_name and real_name != apt_name:
            print(f"  📝 단지명 보정: '{apt_name}' → '{real_name}'")
            apt_name = real_name
            detail["kaptName"] = real_name

    type_label = "아파트" if property_type == "apt" else "오피스텔"
    print(f"\n🏢 건물 유형 확정: {type_label}")

    # 4~6. 병렬 조회: 전월세 + 주변단지 + 학교 + 지하철
    #   (학교/지하철은 독립적이므로 전월세와 동시 실행)
    from concurrent.futures import ThreadPoolExecutor, as_completed

    _rent_result = [None, None]  # [jeonse, wolse]
    _schools_result = [None]
    _stations_result = [None]

    def _fetch_rent_task():
        j, w = fetch_rent_auto(apt_name, lawd_cd, property_type, months=36,
                               road_address=road_address if property_type == "offi" else "",
                               jibun=jibun_str, umd_nm=umd_nm_str)
        _rent_result[0] = j
        _rent_result[1] = w

    def _fetch_schools_task():
        _schools_result[0] = get_nearby_schools(detail["lat"], detail["lon"], radius=3000)

    def _fetch_stations_task():
        _stations_result[0] = get_nearby_stations(detail["lat"], detail["lon"], radius=1000)

    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = [
            executor.submit(_fetch_rent_task),
            executor.submit(_fetch_schools_task),
            executor.submit(_fetch_stations_task),
        ]
        for f in as_completed(futures):
            try:
                f.result()
            except Exception as e:
                print(f"  ⚠️ 병렬 조회 오류: {e}")

    jeonse = _rent_result[0] or []
    wolse  = _rent_result[1] or []
    schools = _schools_result[0] or []
    stations_nearby = _stations_result[0] or []

    # 5. 주변 단지 실거래가
    if property_type == "apt":
        # 아파트: 국토부 단지 목록 기반
        nearby_sales  = fetch_nearby_sales(
            detail["lat"], detail["lon"],
            detail.get("apt_list", []), lawd_cd,
            exclude_name=apt_name, radius=1000, months=36
        )
        nearby_jeonse = fetch_nearby_rent(nearby_sales, lawd_cd, months=36)
    else:
        # 오피스텔: 주변 오피스텔 우선 + 부족분 아파트 보완
        nearby_sales  = build_offi_nearby_sales(
            detail["lat"], detail["lon"], sales, lawd_cd, max_total=5,
            base_name=apt_name
        )
        # 전세 — 주변 단지별 오피스텔 전세 데이터 병렬 조회
        nearby_jeonse = {}
        def _fetch_nearby_rent_one(nm, info):
            j, _ = fetch_rent_auto(nm, lawd_cd, "offi", months=12)
            return nm, j, info

        with ThreadPoolExecutor(max_workers=3) as executor:
            nfutures = [executor.submit(_fetch_nearby_rent_one, nm, info)
                        for nm, info in nearby_sales.items()]
            for f in as_completed(nfutures):
                try:
                    nm, j, info = f.result()
                    if j:
                        nearby_jeonse[nm] = {
                            "trades": j,
                            "dist":   info["dist"],
                            "lat":    info["lat"],
                            "lon":    info["lon"],
                            "marker_type": info.get("marker_type", "offi"),
                        }
                        print(f"  ✅ [전세] {nm}: {len(j)}건")
                except Exception as e:
                    print(f"  ⚠️ 주변 전세 조회 오류: {e}")

    # 7. 공급면적 조회 — DB pyeongs 우선, 없으면 건축물대장 API
    output_dir = os.path.join(output_base, apt_name)
    db_pyeongs = detail.get("_db_pyeongs", [])
    if db_pyeongs:
        supply_map = {round(float(p["exclu"]), 2): round(float(p["supply"]), 2)
                      for p in db_pyeongs if p.get("exclu") and p.get("supply")}
        print(f"  ✅ 공급면적 DB에서 로드: {len(supply_map)}개 평형")
    else:
        supply_map = get_supply_area_map(road_address, property_type=property_type)

    # 8. 이미지 (주변 없으면 단독 라벨 지도로 자동 처리됨)
    img_paths  = generate_images(detail, sales, nearby_sales, nearby_jeonse, output_dir,
                                 auto_mode=auto_mode, supply_map=supply_map, jeonse=jeonse,
                                 property_type=property_type)

    # 9. 원고 생성 전 데이터 검증
    _validate_data(apt_name, sales, jeonse, wolse, supply_map, property_type)

    # 10. 원고 docx
    generate_documents(detail, sales, jeonse, wolse, nearby_sales, schools,
                       output_dir, property_type=property_type, supply_map=supply_map,
                       stations=stations_nearby)

    # 11. 매물 사진 Vision 분석 → 12_매물사진원고.docx
    if photo_paths:
        print(f"\n📷 매물 사진 분석 중 ({len(photo_paths)}장)...")
        try:
            # naver_blog_post.py 임포트
            import importlib.util as _ilu
            _spec2 = _ilu.spec_from_file_location(
                "naver_blog_post",
                os.path.join(os.path.dirname(os.path.abspath(__file__)), "naver_blog_post.py")
            )
            _nbp = _ilu.module_from_spec(_spec2)
            _spec2.loader.exec_module(_nbp)

            photo_results = _nbp.analyze_photos_with_vision(photo_paths, apt_name)

            if photo_results:
                from docx import Document as _Doc
                from docx.shared import Inches as _Inches, Pt as _Pt
                doc = _Doc()
                doc.add_heading(f"{apt_name} 매물 사진", level=1)

                import shutil as _shutil
                photos_dir = os.path.join(output_dir, "매물사진")
                os.makedirs(photos_dir, exist_ok=True)

                for i, pr in enumerate(photo_results, 1):
                    # 공간 구분 헤더
                    doc.add_heading(f"{i}. {pr['space']}", level=2)
                    # 보정 임시파일 → output_dir/매물사진/ 으로 영구 복사
                    ext = os.path.splitext(pr["path"])[1].lower() or ".jpg"
                    saved_photo = os.path.join(photos_dir, f"{i:02d}_{pr['space']}{ext}")
                    try:
                        _shutil.copy2(pr["path"], saved_photo)
                        pic_path = saved_photo
                    except Exception:
                        pic_path = pr["path"]  # 복사 실패 시 원본 경로 fallback
                    # 사진 삽입
                    try:
                        doc.add_picture(pic_path, width=_Inches(5.5))
                    except Exception:
                        doc.add_paragraph(f"[사진: {os.path.basename(pic_path)}]")
                    # 설명 원고
                    if pr.get("caption"):
                        p = doc.add_paragraph(pr["caption"])
                        p.paragraph_format.space_after = _Pt(8)
                save_path = os.path.join(output_dir, "12_매물사진원고.docx")
                doc.save(save_path)
                print(f"  ✅ 매물사진원고 저장: {save_path} ({len(photo_results)}장)")
            else:
                print("  ⚠️ 사진 분석 결과 없음")
        except Exception as e:
            print(f"  ⚠️ 매물사진 원고 생성 오류: {e}")

    print(f"\n{'='*50}")
    print(f"✅ 전체 완료: {output_dir}")
    print(f"  유형:   {type_label}")
    print(f"  이미지: {len([p for p in img_paths.values() if p])}장")
    print(f"  매매:   {len(sales)}건 / 전세: {len(jeonse)}건 / 월세: {len(wolse)}건")
    print(f"  학교:   {len(schools)}개")
    print(f"{'='*50}\n")

    # 메모리 명시적 해제 — 배치 실행 시 단지별 데이터 누적 방지
    del sales, jeonse, wolse, nearby_sales, nearby_jeonse, schools
    import gc; gc.collect()

    return output_dir


# ========================================================
# 실행
# ========================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="휙 단지 조회 + 원고 생성")
    parser.add_argument("input", nargs="?", default=None,
                        help="단지 입력 (예: '중랑구 망우동 효성써너스빌에코 오피스텔 망우로60길 37')")
    parser.add_argument("--auto", action="store_true",
                        help="자동 모드 (라벨지도 확인 건너뜀, 종료 시 엔터 불필요)")
    parser.add_argument("--batch", nargs="+",
                        help="배치 모드: 여러 단지 한번에 처리 (예: --batch '단지1' '단지2')")
    parser.add_argument("--output", default=None,
                        help="출력 디렉토리 (기본값: 바탕화면/원고)")
    args = parser.parse_args()

    if args.batch:
        # 배치 모드 — 에러 나도 다음 단지 계속 진행
        success, fail = [], []
        for i, apt_input in enumerate(args.batch, 1):
            print(f"\n{'#'*50}")
            print(f"# [{i}/{len(args.batch)}] {apt_input}")
            print(f"{'#'*50}")
            try:
                result = run_pipeline(apt_input, output_base=args.output, auto_mode=True)
                if result:
                    success.append(apt_input)
                else:
                    fail.append((apt_input, "데이터 없음"))
            except Exception as e:
                print(f"❌ 오류 발생: {e}")
                fail.append((apt_input, str(e)))

        # 배치 결과 요약
        print(f"\n{'='*50}")
        print(f"📊 배치 결과: 성공 {len(success)}건 / 실패 {len(fail)}건")
        for s in success:
            print(f"  ✅ {s}")
        for f, reason in fail:
            print(f"  ❌ {f} — {reason}")
        print(f"{'='*50}")
    else:
        # 단일 실행 or 기본 배치 리스트
        if args.input:
            run_pipeline(args.input, output_base=args.output, auto_mode=args.auto)
        else:
            print("단지명을 입력하세요.")
            print("예) python test.py '신내동성4차'")
            print("예) python test.py --batch '신내동성4차' '신내우남푸르미아'")

    input("\n종료하려면 엔터...")